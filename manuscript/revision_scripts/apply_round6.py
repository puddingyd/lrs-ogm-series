#!/usr/bin/env python3
"""apply_round6.py — Round-6 patcher: insert Zenodo DOI / ClinVar pending into
Data availability + Code availability statements.

Reads:  manuscript/LRS_OGM_series_2026026_revised_v3.docx
Writes: manuscript/LRS_OGM_series_2026026_revised_v4.docx
"""
import docx, docx.text.paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

INPUT  = 'manuscript/LRS_OGM_series_2026026_revised_v3.docx'
OUTPUT = 'manuscript/LRS_OGM_series_2026026_revised_v4.docx'

ZENODO_DOI = '10.5281/zenodo.19822622'
GITHUB_URL = 'https://github.com/puddingyd/lrs-ogm-series'

NEW_DATA_AVAIL = (
    f'Filtered structural-variant and copy-number-variant summaries '
    f'(per-case manual-review candidate TSVs from the pbsv and HiFiCNV '
    f'pipelines) and consolidated breakpoint coordinates with ACMG/AMP '
    f'classifications for all 14 SVs / CNVs reported in this study are '
    f'available at the project GitHub repository ({GITHUB_URL}) and are '
    f'archived at Zenodo (https://doi.org/{ZENODO_DOI}). The five '
    f'Pathogenic and Likely pathogenic variants identified in the three '
    f'probands have been submitted to ClinVar (accession numbers SCV '
    f'pending; submission date 2026-04-27). Raw HiFi BAM and Bionano '
    f'OGM bnx files contain identifiable sequence-level data and, in '
    f'accordance with the IRB-approved protocol (NCKUH A-BR-109-045), '
    f'are available from the corresponding authors upon reasonable '
    f'request, subject to local ethics-board approval and a data-use '
    f'agreement.'
)
NEW_CODE_AVAIL = (
    f'Custom analysis pipelines accompanying this work — including the '
    f'pbsv and HiFiCNV prioritisation scripts (pbsv_vcf_plus_annotsv_'
    f'summary.sh, hificnv_vcf_plus_annotsv_summary.sh), the BAM-level '
    f'QC script (hifi_bam_qc.sh), the Case 2 retrospective short-read '
    f'WES re-analysis wrappers (Manta, Delly, AUTS2 per-exon coverage), '
    f'and the chromothripsis copy-number-oscillation analysis '
    f'(hificnv_oscillation.py) — are open-source under the MIT licence '
    f'at {GITHUB_URL} and are archived at Zenodo '
    f'(https://doi.org/{ZENODO_DOI}).'
)

doc = docx.Document(INPUT)

def set_para_text(para, new_text):
    """Replace paragraph text via first run; preserve formatting of first run."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(new_text)
    # Force consistent body font (Times New Roman 12)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:cs'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), 'Times New Roman')

# Locate and replace
data_avail_replaced = False
code_avail_replaced = False
for p in doc.paragraphs:
    if p.text.startswith('The data that support the findings'):
        set_para_text(p, NEW_DATA_AVAIL); data_avail_replaced = True
    elif p.text.startswith('Custom analysis scripts are available from'):
        set_para_text(p, NEW_CODE_AVAIL); code_avail_replaced = True

print(f'Data availability replaced: {data_avail_replaced}')
print(f'Code availability replaced: {code_avail_replaced}')

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')

# Verify
chk = docx.Document(OUTPUT)
for p in chk.paragraphs:
    if 'zenodo.19822622' in p.text:
        print(f'\nFOUND in: "{p.text[:80]}..."')
