#!/usr/bin/env python3
"""apply_round5.py
Round-5 patcher (ACMG classification table + main-text additions).

Reads:  LRS_OGM_series_2026026_revised_v2.docx
Writes: LRS_OGM_series_2026026_revised_v3.docx
"""
import re
import docx
import docx.text.paragraph
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

INPUT  = 'LRS_OGM_series_2026026_revised_v2.docx'
OUTPUT = 'LRS_OGM_series_2026026_revised_v3.docx'

doc = docx.Document(INPUT)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def replace_in_para(para, old, new):
    if old not in para.text:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1); return True
    if para.runs:
        para.runs[0].text = para.text.replace(old, new, 1)
        for r in para.runs[1:]:
            r.text = ''
    return True

def replace_all_in_doc(old, new):
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_in_para(p, old, new); n += 1
    return n

def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix): return p
    return None

def find_para_contains(substr):
    for p in doc.paragraphs:
        if substr in p.text: return p
    return None

def insert_para_after(anchor, text='', style='Normal'):
    new_p = OxmlElement('w:p')
    anchor._element.addnext(new_p)
    np = docx.text.paragraph.Paragraph(new_p, anchor._parent)
    try: np.style = doc.styles[style]
    except Exception: pass
    if text:
        np.add_run(text)
    return np

def set_vmerge(cell, val):
    """val: 'restart' (start of merged group) or 'continue' (subsequent rows).

    docx.oxml.OxmlElement('w:vMerge') wraps the element in python-docx's
    CT_VerticalMerge class, which auto-restores w:val='restart' on
    serialization. So we use lxml.etree.SubElement directly to keep the
    element a generic XML node without python-docx's default-attribute logic.
    """
    from lxml import etree
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    tc = cell._tc
    tcPr = tc.find(f'{W}tcPr')
    if tcPr is None:
        tcPr = etree.SubElement(tc, f'{W}tcPr')
        tc.insert(0, tcPr)
    # Remove any existing vMerge
    for vm in tcPr.findall(f'{W}vMerge'):
        tcPr.remove(vm)
    vm = etree.SubElement(tcPr, f'{W}vMerge')
    # python-docx's CT_VerticalMerge subclass auto-adds val='restart' on
    # serialization when val is missing. Force val explicitly in BOTH cases
    # to circumvent this — 'continue' is a valid attribute value per OOXML.
    vm.set(f'{W}val', 'restart' if val == 'restart' else 'continue')

def vmerge_same_value_column(table, col_idx):
    rows = list(table.rows)
    i = 1
    while i < len(rows):
        j = i
        tgt = rows[i].cells[col_idx].text
        while j + 1 < len(rows) and rows[j+1].cells[col_idx].text == tgt:
            j += 1
        if j > i:
            set_vmerge(rows[i].cells[col_idx], 'restart')
            for k in range(i+1, j+1):
                # Clear text first then set vmerge continue
                rows[k].cells[col_idx].text = ''
                set_vmerge(rows[k].cells[col_idx], 'continue')
        i = j + 1

print(f'Loaded: {INPUT}  ({len(doc.paragraphs)} paragraphs)')

# ===========================================================================
# (1) Remove the 1500-fold per-base depth sentence from Discussion Case 2
# ===========================================================================
target_sentence = (
    ' Per-base read depth at the two breakpoints was 0.07× and 0.16×, '
    'against 154.81× at a captured positive-control exon (TP53 exon 4) '
    'and a median of 106.09× across the 19 captured AUTS2 exons '
    '(range 35.90×–274.85×; NM_015570.4) — a more than 1,500-fold deficit '
    'relative to the captured AUTS2 exonic baseline.'
)
# Both with and without the leading space (in case adjacent sentence used different separator)
removed = replace_all_in_doc(target_sentence, '')
if not removed:
    # Try without leading space
    removed = replace_all_in_doc(target_sentence.lstrip(), '')
print(f'  1500-fold sentence removed: {bool(removed)}')

# Also handle the italic-AUTS2 variant of the sentence
target_italic = (
    ' Per-base read depth at the two breakpoints was 0.07× and 0.16×, '
    'against 154.81× at a captured positive-control exon (TP53 exon 4) '
    'and a median of 106.09× across the 19 captured '
)
# Direct text match may fail because AUTS2 might be in italic run. Try searching
# by signature substrings.
for p in doc.paragraphs:
    if 'a more than 1,500-fold deficit' in p.text:
        # Find the run sequence containing this sentence and rebuild paragraph
        full = p.text
        # Locate sentence boundaries
        s_start = full.find('Per-base read depth at the two breakpoints')
        s_end = full.find('captured AUTS2 exonic baseline.', s_start)
        if s_start > 0 and s_end > 0:
            s_end_full = s_end + len('captured AUTS2 exonic baseline.')
            # Backtrack to include preceding space/period
            if s_start > 0 and full[s_start-1] == ' ':
                s_start -= 1
            new_text = full[:s_start] + full[s_end_full:]
            # Rewrite via runs (lose any italics on AUTS2 in this sentence,
            # but we removed the sentence so no concern)
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
            print('  1500-fold sentence stripped via fallback')
        break

# ===========================================================================
# (2) Append ACMG classification sentences to three Discussion paragraphs
# ===========================================================================

def append_to_para(p, text):
    """Append text to the last run of a paragraph (preserves formatting)."""
    if p.runs:
        # Append as a new run with default formatting (avoids inheriting italic
        # if the last run was a gene name)
        # Use plain run with default font
        new_run = p.add_run(text)
        # Try to inherit font from the first non-italic run
        for ref in p.runs[:-1]:
            if not ref.italic and not ref.bold:
                if ref.font.name:
                    new_run.font.name = ref.font.name
                if ref.font.size:
                    new_run.font.size = ref.font.size
                rpr = ref._element.find(qn('w:rPr'))
                if rpr is not None:
                    new_rpr = new_run._element.find(qn('w:rPr'))
                    if new_rpr is not None:
                        new_rpr.getparent().remove(new_rpr)
                    new_run._element.insert(0, deepcopy(rpr))
                break

# Case 1
p_disc_case1 = find_para_contains('REPP-REPD recombination underlies an inversion accompanied by flanking deletions')
if p_disc_case1 is None:
    p_disc_case1 = find_para_contains('In the first case, a de novo inversion on chromosome 8')
if p_disc_case1 is not None:
    append_to_para(p_disc_case1,
        ' The proximal 8p23.1 deletion was classified as Pathogenic, and the '
        'distal 8p22 deletion as Likely pathogenic, under the ACMG/ClinGen 2020 '
        'standards 40 (Supplementary Table 7).')
    print('  Case 1 ACMG sentence appended')

# Case 2
p_disc_case2 = find_para_contains('demonstrating substantial added value for patient care')
if p_disc_case2 is None:
    p_disc_case2 = find_para_contains('The second case carried a de novo inversion on chromosome 7')
if p_disc_case2 is not None:
    append_to_para(p_disc_case2,
        ' Under ACMG/AMP variant interpretation 40, the AUTS2-disrupting '
        'inversion was classified as Pathogenic (Supplementary Table 7).')
    print('  Case 2 ACMG sentence appended')

# Case 3
p_disc_case3 = find_para_contains('CNTNAP2 has the strongest a priori evidence for autism-spectrum')
if p_disc_case3 is None:
    p_disc_case3 = find_para_contains('The third case involved the most complex rearrangement')
if p_disc_case3 is not None:
    append_to_para(p_disc_case3,
        ' Under ACMG/AMP and ACMG/ClinGen 2020 standards 40, the CNTNAP2-disrupting '
        'inversion and the 7q34–q35 deletion were each classified as Likely '
        'pathogenic (Supplementary Table 7).')
    print('  Case 3 ACMG sentence appended')

# ===========================================================================
# (3) Build Supplementary Table 7 — ACMG/AMP classifications (all variants)
# ===========================================================================
# Insert after the Sup Table 6 capture-kit Note (or after Sup Table 6 footnote
# if Note absent).
insertion_anchor = None
for p in doc.paragraphs:
    if p.text.startswith('Note. The original Case 2 WES library'):
        insertion_anchor = p
        break
if insertion_anchor is None:
    for p in doc.paragraphs:
        if p.text.startswith('Footnote.') and 'Manta' in p.text and 'Delly' in p.text:
            insertion_anchor = p
            break
if insertion_anchor is None:
    insertion_anchor = doc.paragraphs[-1]

# Header paragraph
sup7_caption = insert_para_after(insertion_anchor,
    'Supplementary Table 7. ACMG/AMP variant classifications for all '
    'structural and copy-number variants identified in the three probands.',
    style='Normal')

# Build the table data
header = ['Case', 'Variant', 'Coordinates (GRCh38)', 'Size / type',
          'Genes affected', 'ACMG criteria triggered', 'Score',
          'Classification']

rows_data = [
    ('Case 1', '8p23.1 microdeletion (proximal)', 'chr8:8,174,929–10,727,966',
     '2.55 Mb deletion',
     '8p23.1 microdeletion syndrome region; SOX7, TNKS deleted; GATA4 spared',
     '1A; 2A (complete overlap, ClinGen HI sufficient evidence for the recurrent CNV region)',
     '+1.00', 'Pathogenic'),
    ('Case 1', '8p22 deletion (distal)', 'chr8:12,477,756–16,798,965',
     '4.31 Mb deletion',
     '~40 protein-coding genes (incl. TUSC3); no established single HI gene',
     '1A; 3B (25–49 genes); 4O (partial overlap with literature 8p22 NDD cases)',
     '+0.55', 'Likely pathogenic'),
    ('Case 1', '8p inversion (between deletions)',
     'chr8:10,727,966–12,477,756',
     '1.75 Mb inversion (copy-neutral)',
     'No gene disruption identified at breakpoints',
     'None (no HI overlap; mechanism only)', '0.00', 'VUS'),
    ('Case 2', 'AUTS2-disrupting inversion',
     'chr7:69,938,691 ↔ chr7:95,720,897',
     'Copy-neutral inversion; intron 2 of AUTS2 truncated',
     'AUTS2 (ClinGen HI score 3, sufficient evidence for AD haploinsufficiency)',
     'PVS1 (predicted LoF in HI gene); PM6 (de novo by karyotype, not molecularly confirmed); PP4 (phenotype highly specific)',
     '—', 'Pathogenic'),
    ('Case 3', 't(2;13) interchromosomal junction',
     'chr2:190,165,024 ↔ chr13:89,823,743',
     'Translocation breakend',
     'No established disease gene at junctions',
     'None', '0', 'VUS'),
    ('Case 3', '2q33.2–2q32.3 inversion (relocated to 2q33.3)',
     'chr2:195,916,610–204,194,605',
     '8.28 Mb inversion (copy-neutral)',
     'Breakpoints in non-HI region',
     '1A; no HI overlap', '0.00', 'VUS'),
    ('Case 3', '2q33.3 deletion',
     'chr2:208,167,781–208,168,760',
     '1.0 kb deletion',
     'No protein-coding gene (1B)',
     '1B; below size threshold', '0', 'VUS'),
    ('Case 3', '2q33.3–2q34 segment insertion',
     'chr2:208,561,732–212,365,904',
     '3.80 Mb segment relocated to 2q33.3',
     'Copy-neutral relocation; breakpoints in non-HI region',
     '1A; no HI overlap', '0.00', 'VUS'),
    ('Case 3', '2q33.3 local inversion',
     'chr2:208,484,607–208,561,732',
     '77 kb inversion (copy-neutral)',
     'Small; no gene disruption identified',
     '1A; no HI overlap', '0.00', 'VUS'),
    ('Case 3', '13q21.31 fragment translocated to 2q34',
     'chr13:62,374,243–63,090,549 → chr2:212 Mb area',
     '716 kb fragment relocated',
     'Gene-poor region; no established disease gene',
     'None', '0', 'VUS'),
    ('Case 3', 't(5;7) interchromosomal junction',
     'chr5:63,691,635 ↔ chr7:134,079,073',
     'Translocation breakend',
     'No established disease gene at junctions',
     'None', '0', 'VUS'),
    ('Case 3', '7q34–q35 deletion',
     'chr7:139,635,053–145,048,896',
     '5.41 Mb deletion',
     '15 OMIM genes (BRAF, AGK, CLCN1, TBXAS1, WEE2, SSBP1, TAS2R38, PRSS1, PRSS2, TRPV6, KEL, CASP2, NOBOX, TPK1, SLC37A3); ~45 protein-coding',
     '1A; 3B (25–49 genes); 4O (partial overlap with 7q33–q36 NDD literature)',
     '+0.60', 'Likely pathogenic'),
    ('Case 3', '7q35 inversion (CNTNAP2-disrupting)',
     'chr7:145,048,896–146,600,257',
     '1.55 Mb inversion (copy-neutral); breakpoint disrupts CNTNAP2',
     'CNTNAP2 (ClinGen HI score 2; AD risk for ASD/language)',
     'PVS1_moderate (predicted LoF with HI evidence); PM6 (de novo by karyotype, not molecularly confirmed); PP4 (phenotype matches)',
     '—', 'Likely pathogenic'),
    ('Case 3', '4q28.3 deletion',
     'chr4:130,337,827–132,440,009',
     '2.10 Mb deletion',
     'No OMIM disease gene; gene-poor region',
     '1A; 3A (<25 genes); no HI', '0.00', 'VUS, likely benign'),
]

ncols = len(header)
nrows = len(rows_data) + 1
tbl = doc.add_table(rows=nrows, cols=ncols)
tbl.style = 'Table Grid'

# Header
for j, h in enumerate(header):
    c = tbl.rows[0].cells[j]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)

# Body
for i, row in enumerate(rows_data, start=1):
    for j, val in enumerate(row):
        c = tbl.rows[i].cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

# Vertical-merge the Case column (col 0) where consecutive rows share the same case
vmerge_same_value_column(tbl, 0)

# Move the table immediately after the caption paragraph
sup7_caption._element.addnext(tbl._element)

# Footnote
insert_para_after(tbl._element[-1] if False else sup7_caption,
                  '', style='Normal')  # spacer
# Actually find the table to insert footnote after it
# Use the last paragraph in the tbl element's parent's children right after tbl
parent_el = tbl._element.getparent()
tbl_idx = list(parent_el).index(tbl._element)
# Insert a footnote paragraph right after the table
footnote_p = OxmlElement('w:p')
parent_el.insert(tbl_idx + 1, footnote_p)
fp = docx.text.paragraph.Paragraph(footnote_p, sup7_caption._parent)
fp.add_run('Footnote. CNV / SV scores follow the ACMG/ClinGen 2020 technical standards '
           '(Riggs et al., reference 40). Gene-disrupting copy-neutral structural variants '
           '(translocations, inversion breakpoints) were additionally classified under the '
           'ACMG/AMP 2015 framework using PVS1 (predicted loss-of-function in an established '
           'haploinsufficient gene), PS2/PM6 (de novo status; here PM6 since de novo was '
           'inferred from parental karyotyping rather than confirmed at base-pair resolution), '
           'and PP4 (phenotype highly specific for the gene).')
for r in fp.runs:
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')

# Verification
chk = docx.Document(OUTPUT)
checks = [
    ('1500-fold sentence removed',
        lambda d: not any('1,500-fold deficit' in p.text for p in d.paragraphs)),
    ('Case 1 ACMG sentence',
        lambda d: any('proximal 8p23.1 deletion was classified as Pathogenic' in p.text for p in d.paragraphs)),
    ('Case 2 ACMG sentence',
        lambda d: any('AUTS2-disrupting inversion was classified as Pathogenic' in p.text for p in d.paragraphs)),
    ('Case 3 ACMG sentence',
        lambda d: any('each classified as Likely pathogenic' in p.text for p in d.paragraphs)),
    ('Sup Table 7 caption',
        lambda d: any(p.text.startswith('Supplementary Table 7. ACMG') for p in d.paragraphs)),
    ('Sup Table 7 footnote',
        lambda d: any(p.text.startswith('Footnote. CNV / SV scores follow') for p in d.paragraphs)),
    ('Sup Table 7 has 14 data rows',
        lambda d: any(len(t.rows) == 15 and t.rows[0].cells[0].text == 'Case' for t in d.tables)),
]
print()
print('=== Round-5 verification ===')
for name, fn in checks:
    try:
        ok = fn(chk)
        print(f'  [{"OK" if ok else "FAIL"}] {name}')
    except Exception as e:
        print(f'  [ERR {e}] {name}')
