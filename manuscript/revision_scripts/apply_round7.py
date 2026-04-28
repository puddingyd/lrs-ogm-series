#!/usr/bin/env python3
"""apply_round7.py — downgrade 8p22 deletion (NCKUH_LRS_002) from LP to VUS,
and remove SNP array / Read depth from the SV-method field of the two
inversion ClinVar rows (NCKUH_LRS_003, _005). Also update the manuscript
Sup Table 7 row, the Discussion Case 1 sentence, and breakpoint_coordinates.tsv.

Reads:  manuscript/LRS_OGM_series_2026026_revised_v4.docx (user-edited)
        data/clinvar_submissions.xlsx
        data/breakpoint_coordinates.tsv
Writes: manuscript/LRS_OGM_series_2026026_revised_v5.docx
        data/clinvar_submissions.xlsx (in place)
        data/breakpoint_coordinates.tsv (in place)
"""
import re
import openpyxl
import warnings
warnings.filterwarnings('ignore')

# ===========================================================================
# 1. ClinVar Excel — three row updates
# ===========================================================================
XL = 'data/clinvar_submissions.xlsx'
wb = openpyxl.load_workbook(XL, data_only=False)
ws = wb['Variant']

COL = {
    'Local ID': 1, 'Variant type': 11,
    'Germline classification': 34, 'Assertion score': 35,
    'Comment on classification': 37,
    'SV method': 47,
}

# Inversion-only SV-method (no SNP array, no Read depth)
INV_SV_METHOD = 'Paired-end mapping | Sequence alignment | Optical mapping'

# Updated NCKUH_LRS_002 ACMG explanation
NEW_002_COMMENT = (
    'Classified per ACMG/ClinGen 2020 (Riggs et al., reference 40 in '
    'manuscript) for copy-number losses: 1A (protein-coding genes present, '
    '+0.00); 2H (no established haploinsufficient gene, +0.00); 3A '
    '(7 protein-coding RefSeq genes wholly or partially included — '
    'C8orf48, DLC1, LONRF1, MSR1, SGCZ, TRMT9B, TUSC3 — below the 25-gene '
    'threshold; +0.00); 4O (partial overlap with reported 8p22 NDD cases '
    'is weak because most published 8p22 deletions co-occur with 8p23.1 '
    'deletions, confounding phenotype attribution; +0.10). Total score '
    '+0.10 -> Uncertain significance. Note: this deletion co-occurs with '
    'a Pathogenic 8p23.1 microdeletion (NCKUH_LRS_001) and an intervening '
    'copy-neutral inversion in the same proband as part of a single '
    'de novo deletion-inversion-deletion event; an independent contribution '
    'to the proband\'s phenotype cannot be excluded but is not supported '
    'by the formal Riggs scoring.'
)

updated_rows = []
for r in range(6, 50):
    lid = ws.cell(r, COL['Local ID']).value
    if not lid:
        break
    if lid == 'NCKUH_LRS_002':
        ws.cell(r, COL['Germline classification']).value = 'Uncertain significance'
        ws.cell(r, COL['Assertion score']).value = '+0.10'
        ws.cell(r, COL['Comment on classification']).value = NEW_002_COMMENT
        updated_rows.append(f'  Row {r} {lid}: LP -> Uncertain significance; +0.55 -> +0.10; comment rewritten (3A justified by n=7 genes)')
    elif lid in ('NCKUH_LRS_003', 'NCKUH_LRS_005'):
        ws.cell(r, COL['SV method']).value = INV_SV_METHOD
        updated_rows.append(f'  Row {r} {lid}: SV method -> "{INV_SV_METHOD}" (removed SNP array + Read depth — copy-neutral inversion not detectable by either)')
print('=== ClinVar Excel updates ===')
for u in updated_rows:
    print(u)
wb.save(XL)
print(f'Saved: {XL}\n')

# ===========================================================================
# 2. breakpoint_coordinates.tsv — case1_8p22_deletion classification & score
# ===========================================================================
TSV = 'data/breakpoint_coordinates.tsv'
with open(TSV) as f:
    lines = f.readlines()

new_lines = []
n_changed = 0
for line in lines:
    if line.startswith('1\tcase1_8p22_deletion'):
        # Replace last two fields (acmg_classification, acmg_score)
        # Header tail: ... acmg_classification \t acmg_score
        new_line = line.replace('Likely pathogenic', 'VUS').replace('+0.55', '+0.10')
        new_lines.append(new_line)
        n_changed += 1
    else:
        new_lines.append(line)
with open(TSV, 'w') as f:
    f.writelines(new_lines)
print(f'=== breakpoint_coordinates.tsv ===')
print(f'  case1_8p22_deletion row updated: classification "Likely pathogenic" -> "VUS", score +0.55 -> +0.10  ({n_changed} replacement)\n')

# ===========================================================================
# 3. Manuscript v4 -> v5 (preserve user's hand edits to v4)
# ===========================================================================
import docx, docx.text.paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

V4 = 'manuscript/LRS_OGM_series_2026026_revised_v4.docx'
V5 = 'manuscript/LRS_OGM_series_2026026_revised_v5.docx'

doc = docx.Document(V4)

# 3a. Sup Table 7 — find the 8p22 distal deletion row, update cols
sup7 = None
for t in doc.tables:
    if t.rows[0].cells[0].text == 'Case':
        sup7 = t; break

if sup7 is None:
    raise RuntimeError('Sup Table 7 not found')

target_row_idx = None
for ri, row in enumerate(sup7.rows):
    if '8p22 deletion' in row.cells[1].text:
        target_row_idx = ri
        break

if target_row_idx is None:
    raise RuntimeError('8p22 row not found in Sup Table 7')

def set_cell_text(cell, new_text):
    """Replace cell text via first run in first paragraph; preserve formatting."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
        for r in cell.paragraphs[0].runs[1:]:
            r.text = ''
    else:
        cell.text = new_text

row = sup7.rows[target_row_idx]
# Col 4: Genes affected — clarify exact count
set_cell_text(row.cells[4], '7 protein-coding RefSeq genes (C8orf48, DLC1, LONRF1, MSR1, SGCZ, TRMT9B, TUSC3); no established HI gene')
# Col 5: ACMG criteria triggered
set_cell_text(row.cells[5], '1A; 2H (no HI); 3A (7 genes, below 25-gene threshold); 4O (weak — most published 8p22 NDD cases co-occur with 8p23.1)')
# Col 6: Score
set_cell_text(row.cells[6], '+0.10')
# Col 7: Classification
set_cell_text(row.cells[7], 'VUS')
print('=== Manuscript Sup Table 7 ===')
print('  Row "8p22 deletion (distal)": classification LP -> VUS; score +0.55 -> +0.10; criteria 3B -> 3A')

# 3b. Discussion Case 1 paragraph — the ACMG sentence I added in round 5
old_sentence = ('The proximal 8p23.1 deletion was classified as Pathogenic, '
                'and the distal 8p22 deletion as Likely pathogenic, '
                'under the ACMG/ClinGen 2020 standards 40 (Supplementary Table 7).')
new_sentence = ('The proximal 8p23.1 deletion was classified as Pathogenic '
                'under the ACMG/ClinGen 2020 standards 40, while the distal '
                '8p22 deletion meets only Section 3A scoring (7 protein-coding '
                'genes; +0.00) and was classified as a variant of uncertain '
                'significance (Supplementary Table 7); however, because the '
                '8p22 deletion forms part of a single de novo deletion–'
                'inversion–deletion event together with the Pathogenic 8p23.1 '
                'deletion, an independent contribution to the proband\'s '
                'phenotype cannot be excluded.')

replaced = False
for p in doc.paragraphs:
    if old_sentence in p.text:
        # Replace within the paragraph: find the run(s) containing the sentence
        # Simpler: rebuild paragraph text with the substitution
        new_text = p.text.replace(old_sentence, new_sentence)
        # Preserve formatting of first run
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
        # Force consistent body font
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rfonts.set(qn(attr), 'Times New Roman')
        replaced = True
        break

print(f'\n=== Manuscript Discussion Case 1 ===')
print(f'  ACMG sentence replaced: {replaced}')

doc.save(V5)
print(f'\nSaved: {V5}')
