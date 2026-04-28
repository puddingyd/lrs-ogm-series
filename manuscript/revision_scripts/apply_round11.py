#!/usr/bin/env python3
"""apply_round11.py
Cosmetic / scoring fix for the Case 3 4q28.3 deletion (case3_4q28.3_deletion):
the precise UCSC RefSeqCurated count returned 0 protein-coding genes plus
3 non-coding-only entries in chr4:130,337,827-132,440,009 (2.10 Mb).

Riggs 2020 final scoring becomes:
  1B (no protein-coding gene, -0.60) + no HI + no literature  =  -0.60
which is still inside the VUS band (-0.89 to +0.89) but on the negative
side. We tighten the wording to plain "VUS" (matching the 2q33.3 small
deletion already done in round-7).

Reads/writes:
  data/breakpoint_coordinates.tsv (in-place)
  manuscript/LRS_OGM_series_2026026_revised_v6.docx ->
  manuscript/LRS_OGM_series_2026026_revised_v7.docx
"""
import openpyxl, warnings
warnings.filterwarnings('ignore')

# ===========================================================================
# 1. breakpoint_coordinates.tsv — case3_4q28.3_deletion
# ===========================================================================
TSV = 'data/breakpoint_coordinates.tsv'
with open(TSV) as f:
    lines = f.readlines()

new_lines = []
for line in lines:
    if line.startswith('3\tcase3_4q28.3_deletion\t'):
        cols = line.rstrip('\n').split('\t')
        # genes_affected (col 10), acmg_classification (col 12), acmg_score (col 13)
        cols[10] = ('0 protein-coding RefSeq genes (3 non-coding only); '
                    'no OMIM disease gene; gene-poor region')
        cols[12] = 'VUS'
        cols[13] = '-0.60'
        line = '\t'.join(cols) + '\n'
        print('  TSV: case3_4q28.3_deletion updated -> VUS, score -0.60, '
              'genes_affected refined (0 protein-coding RefSeq genes)')
    new_lines.append(line)

with open(TSV, 'w') as f:
    f.writelines(new_lines)

# ===========================================================================
# 2. Manuscript v6 -> v7 (preserve user edits)
# ===========================================================================
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

V6 = 'manuscript/LRS_OGM_series_2026026_revised_v6.docx'
V7 = 'manuscript/LRS_OGM_series_2026026_revised_v7.docx'

doc = docx.Document(V6)

def set_cell_text(cell, new_text):
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
        for r in cell.paragraphs[0].runs[1:]:
            r.text = ''
    else:
        cell.text = new_text

# Find Sup Table 7 4q28.3 row
sup7 = next(t for t in doc.tables if t.rows[0].cells[0].text == 'Case')
for ri, row in enumerate(sup7.rows):
    if '4q28.3 deletion' in row.cells[1].text:
        # Col 4 Genes affected
        set_cell_text(row.cells[4],
            '0 protein-coding RefSeq genes (3 non-coding only); '
            'no OMIM disease gene; gene-poor region')
        # Col 5 ACMG criteria triggered
        set_cell_text(row.cells[5],
            '1B (no protein-coding gene, -0.60); no HI; no literature evidence')
        # Col 6 Score
        set_cell_text(row.cells[6], '-0.60')
        # Col 7 Classification
        set_cell_text(row.cells[7], 'VUS')
        print(f'  Sup Table 7 row "4q28.3 deletion" updated -> VUS, -0.60, 1B (n=0 protein-coding)')
        break

doc.save(V7)
print(f'\nSaved: {V7}')

# Verify
chk = docx.Document(V7)
for t in chk.tables:
    if t.rows[0].cells[0].text != 'Case': continue
    for row in t.rows:
        if '4q28.3 deletion' in row.cells[1].text:
            print('\n=== Verification — Sup Table 7 row "4q28.3 deletion" ===')
            print(f'  Genes affected: {row.cells[4].text}')
            print(f'  ACMG criteria:  {row.cells[5].text}')
            print(f'  Score:          {row.cells[6].text}')
            print(f'  Classification: {row.cells[7].text}')
    break

with open(TSV) as f:
    for line in f:
        if 'case3_4q28.3_deletion' in line:
            cols = line.rstrip().split('\t')
            print('\n=== Verification — breakpoint TSV case3_4q28.3_deletion ===')
            print(f'  genes_affected:      {cols[10]}')
            print(f'  acmg_classification: {cols[12]}')
            print(f'  acmg_score:          {cols[13]}')
