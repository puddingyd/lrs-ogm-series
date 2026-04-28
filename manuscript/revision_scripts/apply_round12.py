#!/usr/bin/env python3
"""apply_round12.py
Round-12 patcher per second independent reviewer.

Reads:  manuscript/LRS_OGM_series_2026026_revised_v7.docx  (user's latest)
        README.md
Writes: manuscript/LRS_OGM_series_2026026_revised_v8.docx
        README.md  (in place)
        deletes  scripts/chromothripsis/hificnv_oscillation.py

Changes:
  Major 1  Add PVS1 vs PVS1_moderate justification (Sup Table 7 footnote)
  Major 2  Remove hificnv_oscillation.py from repo + README mentions
  Major 5  Rewrite Case 1 mechanism wording — single complex event
           (FoSTeS/MMBIR) preferred over two independent events
  Minor 3  Reference list — fix refs 45 (Cameron 2019) + 46 (Belyeu 2021)
           italic journal + bold volume formatting
  Minor 4  Gene italics sweep on Discussion paragraphs
  Minor 6  Data availability date: 2026-04-27 -> 2026-04-28; rephrase
  Minor 7  Sup Table 2 caption: add "after rarity filtering" clarification
  Minor 9  Figure 1D legend: add OGM vs LRS coordinate clarification
  Minor 10 Figure 3D legend: add OGM vs LRS coordinate clarification
"""
import os
import re
import docx
import docx.text.paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

V7 = 'manuscript/LRS_OGM_series_2026026_revised_v7.docx'
V8 = 'manuscript/LRS_OGM_series_2026026_revised_v8.docx'

doc = docx.Document(V7)

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def replace_in_para(para, old, new):
    if old not in para.text:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    if para.runs:
        para.runs[0].text = para.text.replace(old, new, 1)
        for r in para.runs[1:]:
            r.text = ''
    return True

def find_para_starts(prefix):
    return [p for p in doc.paragraphs if p.text.startswith(prefix)]

def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None

# ===========================================================================
# Major 1. PVS1 vs PVS1_moderate justification — append to Sup Table 7 footnote
# ===========================================================================
sup7_footnote = None
for p in doc.paragraphs:
    if p.text.startswith('Footnote. CNV / SV scores follow the ACMG/ClinGen 2020'):
        sup7_footnote = p; break

if sup7_footnote is not None:
    extra = (' The PVS1 strength differs between the two gene-disrupting '
             'copy-neutral inversions (full PVS1 for AUTS2 versus '
             'PVS1_moderate for CNTNAP2) because AUTS2 is rated ClinGen '
             'Dosage Sensitivity haploinsufficiency score 3 (sufficient '
             'evidence for autosomal-dominant haploinsufficiency), whereas '
             'CNTNAP2 is rated haploinsufficiency score 2 (limited evidence '
             'for AD haploinsufficiency in addition to its AR Pitt-Hopkins-'
             'like-1 syndrome), per the ClinGen Sequence Variant '
             'Interpretation Working Group recommendation that PVS1 weight '
             'be modulated by HI evidence strength.')
    if sup7_footnote.runs:
        sup7_footnote.runs[-1].text = sup7_footnote.runs[-1].text + extra
    print('Major 1: appended PVS1 vs PVS1_moderate justification to Sup Table 7 footnote')

# ===========================================================================
# Major 5. Rewrite Case 1 mechanism (parsimony) in Discussion
# ===========================================================================
old5 = ('In terms of mechanism, the breakpoint anatomy pointed to a two-step '
        'rearrangement process. The inversion breakpoint on 8p mapped to the '
        'low-copy repeats REPP and REPD, supporting NAHR as the initiating '
        'mechanism, whereas 4-bp microhomology at the distal deletion '
        'junctions was consistent with MMEJ. Taken together, these findings '
        'are compatible with a two-step model in which an NAHR event '
        'between REPP and REPD generated the inversion and subsequent '
        'error-prone repair (MMEJ, or alternatively replication-based '
        'mechanisms such as FoSTeS / MMBIR) yielded the flanking deletions.')
new5 = ('In terms of mechanism, the breakpoint anatomy carries signatures '
        'of distinct repair classes at different junctions: the inversion '
        'breakpoints fall within the REPP / REPD low-copy-repeat pair, a '
        'configuration commonly mediated by non-allelic homologous '
        'recombination (NAHR), whereas the flanking-deletion junctions '
        'display 4-bp microhomology consistent with end-joining-class '
        'repair. Rather than two independent de novo events, these features '
        'are most parsimoniously explained by a single complex repair '
        'episode — for example replication-based template switching (FoSTeS '
        '/ MMBIR) — in which different junction classes within the same '
        'event produce the apparent NAHR-like and microhomology-mediated '
        'signatures.')

found5 = False
for p in doc.paragraphs:
    if old5 in p.text:
        new_text = p.text.replace(old5, new5)
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rfonts.set(qn(attr), 'Times New Roman')
        found5 = True; break
print(f'Major 5: Case 1 mechanism wording rewritten: {found5}')

# ===========================================================================
# Minor 3. Refs 45 + 46 — italic journal + bold volume per Nature style
# ===========================================================================
def reformat_nature_ref(p, journal, volume, before, after):
    """Rebuild paragraph as: '<before>' + italic '<journal>' + ' ' +
       bold '<volume>' + '<after>'."""
    # Clear all runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r1 = p.add_run(before)
    r2 = p.add_run(journal); r2.italic = True
    r3 = p.add_run(' ')
    r4 = p.add_run(volume); r4.bold = True
    r5 = p.add_run(after)

for p in doc.paragraphs:
    if p.text.startswith('45. Cameron'):
        # 45. Cameron, D.L., Di Stefano, L. & Papenfuss, A.T. Comprehensive
        # evaluation and characterisation of short read general-purpose
        # structural variant calling software. Nat Commun 10, 3240 (2019).
        before  = ('45. Cameron, D.L., Di Stefano, L. & Papenfuss, A.T. '
                   'Comprehensive evaluation and characterisation of short '
                   'read general-purpose structural variant calling software. ')
        journal = 'Nat Commun'
        volume  = '10'
        after   = ', 3240 (2019).'
        reformat_nature_ref(p, journal, volume, before, after)
        print('Minor 3: ref 45 (Cameron 2019) reformatted to Nature style')
    elif p.text.startswith('46. Belyeu'):
        before  = ('46. Belyeu, J.R., Brand, H., Wang, H., Zhao, X., '
                   'Pedersen, B.S., Layer, R.M. et al. De novo structural '
                   'mutation rates and gamete-of-origin biases revealed '
                   'through genome sequencing of 2,396 families. ')
        journal = 'Am J Hum Genet'
        volume  = '108'
        after   = ', 597–607 (2021).'
        reformat_nature_ref(p, journal, volume, before, after)
        print('Minor 3: ref 46 (Belyeu 2021) reformatted to Nature style')

# ===========================================================================
# Minor 4. Gene italics sweep — body paragraphs
# ===========================================================================
GENES = ['AUTS2', 'CNTNAP2', 'GATA4', 'SOX7', 'TNKS', 'MCPH1', 'BRAF',
         'AGK', 'CLCN1', 'TBXAS1', 'WEE2', 'SSBP1', 'TAS2R38', 'PRSS1',
         'PRSS2', 'TRPV6', 'KEL', 'CASP2', 'NOBOX', 'TPK1', 'SLC37A3',
         'FMR1', 'TP53', 'TUSC3', 'DLC1', 'MSR1', 'SGCZ', 'LONRF1',
         'TRMT9B', 'C8orf48']
gene_re = re.compile(r'\b(' + '|'.join(re.escape(g) for g in GENES) + r')\b')

n_italicised = 0
for p in doc.paragraphs:
    # Skip references and figure captions to avoid touching their formatting
    if re.match(r'^\d+\.\s', p.text):
        continue
    if p.text.startswith('Figure '):
        continue
    if (len(p.text) > 2 and p.text[1:3] == '. ' and p.text[0].isalpha()):
        continue  # figure panel "A. ", "B. ", etc.
    for run in list(p.runs):
        text = run.text
        if not text or run.italic:
            continue
        if not gene_re.search(text):
            continue
        rpr = run._element.find(qn('w:rPr'))
        parent = run._element.getparent()
        idx = list(parent).index(run._element)
        parent.remove(run._element)
        parts = []
        last = 0
        for m in gene_re.finditer(text):
            if m.start() > last:
                parts.append(('plain', text[last:m.start()]))
            parts.append(('italic', m.group()))
            last = m.end()
        if last < len(text):
            parts.append(('plain', text[last:]))
        for kind, t in parts:
            new_run = p.add_run(t)
            if rpr is not None:
                new_rpr = new_run._element.find(qn('w:rPr'))
                if new_rpr is not None:
                    new_rpr.getparent().remove(new_rpr)
                new_run._element.insert(0, deepcopy(rpr))
            if kind == 'italic':
                new_run.italic = True
            new_run._element.getparent().remove(new_run._element)
            parent.insert(idx, new_run._element)
            idx += 1
        n_italicised += 1
print(f'Minor 4: gene-italics sweep — {n_italicised} runs touched')

# ===========================================================================
# Minor 6. Data availability date 2026-04-27 -> "submitted on 2026-04-28, accessions pending"
# ===========================================================================
n6 = 0
for p in doc.paragraphs:
    if 'submission date 2026-04-27' in p.text:
        if replace_in_para(p, 'SCV pending; submission date 2026-04-27',
                              'submitted on 2026-04-28, accessions pending'):
            n6 += 1
print(f'Minor 6: Data availability date updated -> "submitted on 2026-04-28, accessions pending" ({n6} replacement)')

# ===========================================================================
# Minor 7. Sup Table 2 caption — add "after rarity filtering" clarification
# ===========================================================================
n7 = 0
for p in doc.paragraphs:
    if p.text.startswith('Supplementary Table 2. Filtering strategy and number of OGM variant calls'):
        # Append a clarifying clause
        new_text = p.text.rstrip('.').rstrip() + ' (counts shown are after rarity filtering against the Bionano control database).'
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
        n7 += 1
print(f'Minor 7: Sup Table 2 caption clarified ({n7} replacement)')

# ===========================================================================
# Minor 9. Figure 1D legend — add OGM-vs-LRS coordinate clarification
# ===========================================================================
n9 = 0
for p in doc.paragraphs:
    if p.text.startswith('D.') and 'OGM evidence of intrachromosomal recombination between 8p23.1' in p.text:
        # Append at the end of panel D legend
        suffix = (' The OGM-defined junction (chr8:10,756,153) lies within '
                  'the chr8:10,727,966–10,756,153 OGM-label interval that '
                  'flanks the LRS-refined breakpoint (chr8:10,727,966); '
                  'OGM and LRS coordinates therefore differ by ~28 kb '
                  'reflecting OGM label resolution versus base-pair LRS '
                  'resolution.')
        if p.runs:
            p.runs[-1].text = p.runs[-1].text + suffix
        n9 += 1; break
print(f'Minor 9: Figure 1D legend coordinate clarification appended ({n9} replacement)')

# ===========================================================================
# Minor 10. Figure 3D legend — add OGM-vs-LRS coordinate clarification
# ===========================================================================
n10 = 0
for p in doc.paragraphs:
    if p.text.startswith('D.') and 'OGM evidence of an intrachromosomal fusion between 7q34' in p.text:
        suffix = (' The OGM-defined junction (chr7:139,632,423) and the '
                  'LRS-refined deletion start (chr7:139,635,053) differ by '
                  '~2.6 kb reflecting OGM label resolution versus base-pair '
                  'LRS resolution.')
        if p.runs:
            p.runs[-1].text = p.runs[-1].text + suffix
        n10 += 1; break
print(f'Minor 10: Figure 3D legend coordinate clarification appended ({n10} replacement)')

doc.save(V8)
print(f'\nSaved: {V8}')

# ===========================================================================
# Major 2. Remove hificnv_oscillation.py + tidy README mentions
# ===========================================================================
osc_path = 'scripts/chromothripsis/hificnv_oscillation.py'
if os.path.exists(osc_path):
    os.remove(osc_path)
    print(f'\nMajor 2: removed {osc_path}')
osc_dir = 'scripts/chromothripsis'
if os.path.isdir(osc_dir) and not os.listdir(osc_dir):
    os.rmdir(osc_dir)
    print(f'           removed empty directory {osc_dir}')

# README cleanup
README = 'README.md'
if os.path.exists(README):
    with open(README) as f:
        text = f.read()
    new_text = text
    # Remove the chromothripsis tree node + its line
    new_text = re.sub(r'^.*chromothripsis/.*$\n', '', new_text, flags=re.M)
    new_text = re.sub(r'^.*hificnv_oscillation\.py.*$\n', '', new_text, flags=re.M)
    # Remove the Section 3 chromothripsis subsection (heading + commands)
    new_text = re.sub(
        r'### 3\. Chromothripsis assessment.*?(?=^### |\Z)',
        '', new_text, count=1, flags=re.M | re.S)
    # Remove "the chromothripsis copy-number-oscillation analysis"-like phrases in availability paragraphs
    new_text = new_text.replace(
        ', and the chromothripsis copy-number-oscillation analysis '
        '(`hificnv_oscillation.py`)',
        '')
    if new_text != text:
        with open(README, 'w') as f:
            f.write(new_text)
        print('Major 2: README mentions of hificnv_oscillation removed')
