#!/usr/bin/env python3
"""apply_round10.py
Upgrade NCKUH_LRS_004 (Case 3 7q34-q35 deletion) from Likely pathogenic
to Pathogenic after a precise UCSC RefSeq gene count gave n=69 (3C, +0.90)
rather than the n=25-49 / 3B (+0.45) I had assumed.

Riggs 2020 final scoring:
  1A (+0.00)  +  3C (69 genes, +0.90)  +  4O (+0.15)  =  +1.05  ->  Pathogenic

Reads/writes:
  data/clinvar_submissions.xlsx       (in-place; row 9)
  data/breakpoint_coordinates.tsv     (in-place)
  manuscript/LRS_OGM_series_2026026_revised_v5.docx ->
  manuscript/LRS_OGM_series_2026026_revised_v6.docx
"""
import re
import openpyxl, warnings
warnings.filterwarnings('ignore')

# ===========================================================================
# 1. ClinVar Excel — row 9 (NCKUH_LRS_004)
# ===========================================================================
XL = 'data/clinvar_submissions.xlsx'
wb = openpyxl.load_workbook(XL, data_only=False)
ws = wb['Variant']

NEW_COMMENT = (
    'Classified per ACMG/ClinGen 2020 (Riggs et al., reference 40 in '
    'manuscript) for copy-number losses: 1A (protein-coding genes '
    'present, +0.00); 3C (69 protein-coding RefSeq genes wholly or '
    'partially included, well above the 35-gene threshold; +0.90); 4O '
    '(partial overlap with reported 7q33-q36 NDD cases — Dilzell 2015 '
    'PMID:26064708, Kale 2016 PMID:28053794; +0.15). Total score +1.05 '
    '-> Pathogenic.'
)

for r in range(6, 50):
    if ws.cell(r, 1).value == 'NCKUH_LRS_004':
        ws.cell(r, 34).value = 'Pathogenic'        # Germline classification
        ws.cell(r, 37).value = NEW_COMMENT          # Comment on classification
        # col 35 (Assertion score) was cleared in round-9; leave empty
        print(f'  ClinVar row {r} NCKUH_LRS_004: classification LP -> Pathogenic; comment rewritten (3C with n=69 genes)')
        break

wb.save(XL)
print(f'  Saved: {XL}\n')

# ===========================================================================
# 2. breakpoint_coordinates.tsv — case3_7q34_q35_deletion row
# ===========================================================================
TSV = 'data/breakpoint_coordinates.tsv'
with open(TSV) as f:
    text = f.read()

old_genes = ('15 OMIM genes (BRAF, AGK, CLCN1, TBXAS1, WEE2, SSBP1, TAS2R38, '
             'PRSS1, PRSS2, TRPV6, KEL, CASP2, NOBOX, TPK1, SLC37A3); '
             '~45 protein-coding')
new_genes = ('69 protein-coding RefSeq genes (15 OMIM-listed: BRAF, AGK, '
             'CLCN1, TBXAS1, WEE2, SSBP1, TAS2R38, PRSS1, PRSS2, TRPV6, '
             'KEL, CASP2, NOBOX, TPK1, SLC37A3)')

if old_genes in text:
    text = text.replace(old_genes, new_genes)
    print(f'  TSV: genes_affected updated for case3_7q34_q35_deletion ({len(new_genes)} chars)')

# Replace classification + score on the same row
for old_seg, new_seg in [
    ('case3_7q34_q35_deletion\tchr7\t139635053\t145048896\tchr7\t139635053\t145048896\t5413843\tdeletion',
     None),  # row anchor; we just need the line
]:
    pass

# Simple line-based replace: find the row and rewrite the last two columns
new_lines = []
fixed = False
for line in text.splitlines(keepends=True):
    if line.startswith('3\tcase3_7q34_q35_deletion\t'):
        line = line.replace('Likely pathogenic', 'Pathogenic')
        line = line.replace('+0.60', '+1.05')
        fixed = True
    new_lines.append(line)

with open(TSV, 'w') as f:
    f.writelines(new_lines)
print(f'  TSV: case3_7q34_q35_deletion classification "Likely pathogenic" -> "Pathogenic"; score +0.60 -> +1.05  (fixed={fixed})\n')

# ===========================================================================
# 3. Manuscript v5 -> v6 (preserve any user edits to v5)
# ===========================================================================
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from copy import deepcopy

V5 = 'manuscript/LRS_OGM_series_2026026_revised_v5.docx'
V6 = 'manuscript/LRS_OGM_series_2026026_revised_v6.docx'

doc = docx.Document(V5)

# 3a. Sup Table 7 — find "7q34-q35 deletion" row and update
sup7 = next(t for t in doc.tables if t.rows[0].cells[0].text == 'Case')

def set_cell_text(cell, new_text):
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
        for r in cell.paragraphs[0].runs[1:]:
            r.text = ''
    else:
        cell.text = new_text

for ri, row in enumerate(sup7.rows):
    if '7q34' in row.cells[1].text and 'deletion' in row.cells[1].text:
        # Col 4 Genes affected
        set_cell_text(row.cells[4],
            '69 protein-coding RefSeq genes; 15 OMIM-listed (BRAF, AGK, '
            'CLCN1, TBXAS1, WEE2, SSBP1, TAS2R38, PRSS1, PRSS2, TRPV6, '
            'KEL, CASP2, NOBOX, TPK1, SLC37A3)')
        # Col 5 ACMG criteria triggered
        set_cell_text(row.cells[5],
            '1A; 3C (69 genes, well above 35-gene threshold); 4O '
            '(partial overlap with 7q33–q36 NDD literature)')
        # Col 6 Score
        set_cell_text(row.cells[6], '+1.05')
        # Col 7 Classification
        set_cell_text(row.cells[7], 'Pathogenic')
        print(f'  Sup Table 7 row "{row.cells[1].text}" updated -> Pathogenic, +1.05, 3C')
        break

# 3b. Discussion Case 3 ACMG sentence
old_sentence = ('Under ACMG/AMP and ACMG/ClinGen 2020 standards 40, the '
                'CNTNAP2-disrupting inversion and the 7q34–q35 deletion '
                'were each classified as Likely pathogenic '
                '(Supplementary Table 7).')
new_sentence = ('Under ACMG/AMP and ACMG/ClinGen 2020 standards 40, the '
                '7q34–q35 deletion was classified as Pathogenic and the '
                'CNTNAP2-disrupting inversion as Likely pathogenic '
                '(Supplementary Table 7).')

replaced = False
for p in doc.paragraphs:
    if old_sentence in p.text:
        new_text = p.text.replace(old_sentence, new_sentence)
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ''
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rfonts.set(qn(attr), 'Times New Roman')
        replaced = True
        break
print(f'  Discussion Case 3 ACMG sentence updated: {replaced}')

doc.save(V6)
print(f'\nSaved: {V6}')
