#!/usr/bin/env python3
"""apply_round4.py
Apply round-4 (independent reviewer) refinements directly to the user's
hand-edited round-3 docx, preserving all of their manual changes.

Reads:  LRS_OGM_series_2026026_revised.docx   (user's edited version)
Writes: LRS_OGM_series_2026026_revised_v2.docx
"""
import re
import docx
import docx.text.paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

INPUT  = 'LRS_OGM_series_2026026_revised.docx'
OUTPUT = 'LRS_OGM_series_2026026_revised_v2.docx'

doc = docx.Document(INPUT)

# ---------------------------------------------------------------------------
# Helpers (same as in revise_manuscript.py)
# ---------------------------------------------------------------------------
def replace_in_para(para, old, new):
    if old not in para.text:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    if para.runs:
        para.runs[0].text = para.text.replace(old, new, 1)
        for r in para.runs[1:]:
            r.text = ''
    return True

def replace_all_in_doc(old, new):
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_in_para(p, old, new); n += 1
    return n

def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None

def insert_para_after(anchor, text='', style='Normal'):
    new_p = OxmlElement('w:p')
    anchor._element.addnext(new_p)
    np = docx.text.paragraph.Paragraph(new_p, anchor._parent)
    try:
        np.style = doc.styles[style]
    except Exception:
        pass
    if text:
        np.add_run(text)
    return np

print(f'Loaded: {INPUT}  ({len(doc.paragraphs)} paragraphs)')

# ===========================================================================
# A items
# ===========================================================================

# (A1+A2) Strip stray markdown markers `*` from references and figure callouts
for p in doc.paragraphs:
    if not any('*' in r.text for r in p.runs if r.text):
        continue
    # Apply to references and figure-callout paragraphs
    if re.match(r'^\d+\.\s', p.text) or 'Figure' in p.text:
        for r in p.runs:
            if r.text and '*' in r.text:
                r.text = r.text.replace('*', '')

# (A3) Limitation paragraph numbering
p_lim = find_para('Our study has several limitations.')
if p_lim:
    replace_in_para(p_lim, 'Finally,', 'Fourth,')
    replace_in_para(p_lim, ' Fifth, ', ' Finally, ')

# (A4) Methods §2.6.2 empty parens
replace_all_in_doc('the PacBio WGS Variant Pipeline ()',
    'the PacBio WGS Variant Pipeline (https://github.com/PacificBiosciences/HiFi-human-WGS-WDL, v2.0.0)')
replace_all_in_doc('the PacBio WGS Variant Pipeline (https://github.com/PacificBiosciences/HiFi-human-WGS-WDL).',
    'the PacBio WGS Variant Pipeline (https://github.com/PacificBiosciences/HiFi-human-WGS-WDL, v2.0.0).')

# (A5) Capture kit Note in Supplementary (after Sup Table 6 footnote)
p_sup6_foot = None
for p in doc.paragraphs:
    if p.text.startswith('Footnote.') and 'Manta' in p.text and 'Delly' in p.text:
        p_sup6_foot = p; break
if p_sup6_foot is not None and not any('Roche KAPA HyperExome Plus' in p.text for p in doc.paragraphs):
    note = ('Note. The original Case 2 WES library was prepared with the Roche '
            'KAPA HyperExome Plus capture kit (Roche Sequencing Solutions). The '
            'AUTS2 intron 2 breakpoint (chr7:69,938,691) and the 7q21.3 partner '
            'breakpoint (chr7:95,720,897) both lie outside the kit\'s capture '
            'target intervals, which explains the near-zero local read depth and '
            'the absence of any short-read SV-caller signal at these positions.')
    insert_para_after(p_sup6_foot, note, style='Normal')

# (A6) Gene italics sweep — body paragraphs only (skip refs and figure legends)
GENES = ['AUTS2','CNTNAP2','GATA4','SOX7','TNKS','MCPH1','BRAF','AGK','CLCN1',
         'TBXAS1','WEE2','SSBP1','TAS2R38','PRSS1','PRSS2','TRPV6','KEL',
         'CASP2','NOBOX','TPK1','SLC37A3','FMR1','TP53']
gene_re = re.compile(r'\b(' + '|'.join(re.escape(g) for g in GENES) + r')\b')

def italicize_gene_runs(p):
    if re.match(r'^\d+\.\s', p.text): return        # references
    if p.text.startswith('Figure '):    return        # figure title
    if len(p.text) > 2 and p.text[1:3] == '. ' and p.text[0].isalpha():
        return                                        # figure panel
    for run in list(p.runs):
        text = run.text
        if not text or run.italic:
            continue
        if not gene_re.search(text):
            continue
        rpr = run._element.find(qn('w:rPr'))
        parent = run._element.getparent()
        idx = list(parent).index(run._element)
        parent.remove(run._element)
        parts = []
        last = 0
        for m in gene_re.finditer(text):
            if m.start() > last:
                parts.append(('plain', text[last:m.start()]))
            parts.append(('italic', m.group()))
            last = m.end()
        if last < len(text):
            parts.append(('plain', text[last:]))
        for kind, t in parts:
            new_run = p.add_run(t)
            if rpr is not None:
                new_rpr = new_run._element.find(qn('w:rPr'))
                if new_rpr is not None:
                    new_rpr.getparent().remove(new_rpr)
                new_run._element.insert(0, deepcopy(rpr))
            if kind == 'italic':
                new_run.italic = True
            new_run._element.getparent().remove(new_run._element)
            parent.insert(idx, new_run._element)
            idx += 1

for p in doc.paragraphs:
    italicize_gene_runs(p)

# ===========================================================================
# B items
# ===========================================================================

# (B1) Case 1 mechanism softening
replace_all_in_doc(
    'This configuration supports non-allelic homologous recombination (NAHR) mediated by REPD and REPP as the source of the inversion.',
    'This configuration is compatible with non-allelic homologous recombination (NAHR) between REPD and REPP as the initiating event of the inversion.')
replace_all_in_doc(
    'support a model in which the inversion was initiated by NAHR between REPP and REPD, with subsequent error-prone repair leading to deletion of the intervening segments through MMEJ.',
    'are compatible with a two-step model in which an NAHR event between REPP and REPD generated the inversion and subsequent error-prone repair (MMEJ, or alternatively replication-based mechanisms such as FoSTeS / MMBIR) yielded the flanking deletions.')

# (B2) Case 2 Figure 2 title + panel F
replace_all_in_doc(
    'Figure 2. Chromosome 7 paracentric inversion of Case 2 resolved by long-read sequencing (LRS) and optical genome mapping (OGM) with proposed MIR3-mediated mechanism',
    'Figure 2. Chromosome 7 paracentric inversion of Case 2 resolved by long-read sequencing (LRS) and optical genome mapping (OGM); proposed non-homologous end joining (NHEJ) mechanism with nearby MIR3 SINE elements.')
replace_all_in_doc(
    'F. Proposed mechanism. Two Short Interspersed Nuclear Element (SINE) fragments annotated by RepeatMasker as MIR3 (Mammalian-wide Interspersed Repeat family 3), one located at 5,324 bp proximal to breakpoint A and the other at 2,274 bp distal to breakpoint D (red boxes; arrowheads indicate orientation), may have facilitated recombination. Insets show the junctional sequences, which share only a single-base (T) homology at the breakpoints.',
    'F. Proposed mechanism. The single-base (T) microhomology at the junctions (insets) is most consistent with non-homologous end joining (NHEJ). Two Short Interspersed Nuclear Element (SINE) fragments annotated by RepeatMasker as MIR3 (Mammalian-wide Interspersed Repeat family 3) lie 5,324 bp proximal to breakpoint A and 2,274 bp distal to breakpoint D (red boxes; arrowheads indicate orientation); these elements do not span the junctions but may have predisposed the locus to ectopic strand exchange.')

# (B3) Case 3 chromoanagenesis harmonisation
replace_all_in_doc(
    'LRS and OGM resolved a complex multi-chromosomal rearrangement reminiscent of chromoanagenesis',
    'LRS and OGM resolved a complex multi-chromosomal genomic rearrangement (CGR) with features suggestive of chromoanagenesis')
replace_all_in_doc(
    'is highly reminiscent of chromoanagenesis — a class of complex genomic rearrangement (CGR) that includes chromothripsis and chromoplexy',
    'shows features suggestive of chromoanagenesis (a class of complex genomic rearrangement, CGR, that includes chromothripsis and chromoplexy)')

# (B4) Case 3 7q deletion attribution
replace_all_in_doc(
    'CNTNAP2 truncation and a cryptic 5.4-Mb 7q34–q35 deletion together explained his phenotype',
    'CNTNAP2 truncation and a cryptic 5.4-Mb 7q34–q35 deletion together contribute to his phenotype')
replace_all_in_doc(
    'CNTNAP2 disruption together with a 5.4-Mb 7q34–q35 deletion plausibly explained his phenotype',
    'CNTNAP2 disruption together with a 5.4-Mb 7q34–q35 deletion together contribute to his phenotype')

# (B5) Manta/Delly continuum + new refs
replace_all_in_doc(
    'demonstrates that the AUTS2 intron 2 breakpoint and its 7q21.3 partner lie outside exome capture territory and are inaccessible to short-read WES at the read level, regardless of the variant caller applied.',
    'indicates that the AUTS2 intron 2 breakpoint and its 7q21.3 partner lie at the assay-sensitivity end of a continuum: short-read WES has both limited capture coverage at intronic positions and reduced caller sensitivity for inversions in repeat-rich contexts 45,46, so neither off-target reads nor a more permissive caller can reliably recover this event.')

# (B6) De novo limitation
replace_all_in_doc(
    'Future cohorts incorporating routine trio confirmation would further strengthen inference of de novo origin and would exclude rare instances of low-level parental mosaicism.',
    'In each proband, de novo status is therefore based on parental karyotyping (resolution of several megabases) and was not molecularly confirmed at base-pair resolution; future cohorts incorporating routine trio LRS or junction-spanning Sanger confirmation would strengthen the inference of de novo origin and exclude rare instances of low-level parental mosaicism.')

# Add refs 45/46 after the last existing ref
last_ref = None
for p in doc.paragraphs:
    if p.text.startswith('44.') and 'Rausch' in p.text:
        last_ref = p
if last_ref is not None and not any(p.text.startswith('45. Cameron') for p in doc.paragraphs):
    new_refs = [
        '45. Cameron, D.L., Di Stefano, L. & Papenfuss, A.T. Comprehensive evaluation and characterisation of short read general-purpose structural variant calling software. Nat Commun 10, 3240 (2019).',
        '46. Belyeu, J.R., Brand, H., Wang, H., Zhao, X., Pedersen, B.S., Layer, R.M. et al. De novo structural mutation rates and gamete-of-origin biases revealed through genome sequencing of 2,396 families. Am J Hum Genet 108, 597–607 (2021).',
    ]
    anchor = last_ref
    for txt in new_refs:
        anchor = insert_para_after(anchor, txt, style='Normal')

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')

# Verification
chk = docx.Document(OUTPUT)
for label, fn in [
    ('PacBio pipeline URL filled',
        lambda d: any('HiFi-human-WGS-WDL, v2.0.0' in p.text for p in d.paragraphs)),
    ('Capture-kit Note added',
        lambda d: any('Roche KAPA HyperExome Plus' in p.text for p in d.paragraphs)),
    ('Case 1 mechanism softened',
        lambda d: any('compatible with non-allelic homologous recombination' in p.text for p in d.paragraphs)),
    ('Fig 2 title NHEJ',
        lambda d: any('proposed non-homologous end joining (NHEJ) mechanism with nearby MIR3' in p.text for p in d.paragraphs)),
    ('Abstract CGR + chromoanagenesis',
        lambda d: any('CGR) with features suggestive of chromoanagenesis' in p.text for p in d.paragraphs)),
    ('"together contribute to"',
        lambda d: any('together contribute to his phenotype' in p.text for p in d.paragraphs)),
    ('Manta continuum + ref 45,46',
        lambda d: any('45,46' in p.text for p in d.paragraphs)),
    ('De novo bp-level caveat',
        lambda d: any('not molecularly confirmed at base-pair resolution' in p.text for p in d.paragraphs)),
    ('Ref 45 Cameron',
        lambda d: any(p.text.startswith('45. Cameron') for p in d.paragraphs)),
    ('Ref 46 Belyeu',
        lambda d: any(p.text.startswith('46. Belyeu') for p in d.paragraphs)),
]:
    ok = fn(chk)
    print(f'  [{"OK" if ok else "FAIL"}] {label}')
