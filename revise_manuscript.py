#!/usr/bin/env python3
"""revise_manuscript.py — reviewer-driven revision of the LRS+OGM manuscript.

Reads:  LRS_OGM series 2026026.docx
Writes: LRS_OGM_series_2026026_revised.docx
"""
import re
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.text.paragraph

INPUT = 'LRS_OGM series 2026026.docx'
OUTPUT = 'LRS_OGM_series_2026026_revised.docx'

doc = docx.Document(INPUT)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def replace_in_para(para, old, new, n=1):
    """Replace occurrences of old in para. Edits within a single run when
    possible to preserve italics/bold; otherwise rewrites first run text."""
    if old not in para.text:
        return False
    # Single-run match
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    # Cross-run: rebuild
    if para.runs:
        new_text = para.text.replace(old, new, n)
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    return True

def replace_all_in_doc(old, new):
    """Replace old → new in every paragraph."""
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_in_para(p, old, new)
            count += 1
    return count

def find_para(prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None

def find_para_contains(substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    return None

def insert_para_after(anchor, text='', style='Normal'):
    new_p = OxmlElement('w:p')
    anchor._element.addnext(new_p)
    np = docx.text.paragraph.Paragraph(new_p, anchor._parent)
    try:
        np.style = doc.styles[style]
    except Exception:
        pass
    if text:
        np.add_run(text)
    return np

def insert_para_before(anchor, text='', style='Normal'):
    new_p = OxmlElement('w:p')
    anchor._element.addprevious(new_p)
    np = docx.text.paragraph.Paragraph(new_p, anchor._parent)
    try:
        np.style = doc.styles[style]
    except Exception:
        pass
    if text:
        np.add_run(text)
    return np

def insert_paras_after(anchor, items):
    """Insert multiple paragraphs in order after anchor.
    items is list of (text, style)."""
    last = anchor
    for text, style in items:
        last = insert_para_after(last, text, style)
    return last

print('Loaded:', INPUT, '|', len(doc.paragraphs), 'paragraphs')

# ===========================================================================
# Step 1. Simple text fixes (typos, single-token edits)
# ===========================================================================
# UMWH → UHMW
replace_all_in_doc('UMWH gDNA', 'UHMW gDNA')

# DLE-1 motif palindromic note
replace_all_in_doc(
    'specifically targeting the sequence motif CTTAAG within the DNA',
    'specifically targeting the palindromic sequence motif CTTAAG (recognized on both strands of the DNA)'
)

# "tract" → "track" in figure legends
replace_all_in_doc('upper tract', 'upper track')
replace_all_in_doc('middle tract', 'middle track')
replace_all_in_doc('lower tract', 'lower track')
replace_all_in_doc('(upper tract)', '(upper track)')

# Figure 3D coordinate fix (chr7:69.9 Mb is wrong, should be chr7:139 Mb)
replace_all_in_doc('chr7:69.9 Mb (upper track)', 'chr7:139 Mb (upper track)')

# Discussion: capital "Whereas" inside sentence
replace_all_in_doc('Moreover, Whereas rearrangements', 'Moreover, whereas rearrangements')

# Section numbering: 3.3 Case 2 → 3.2 Case 2
p = find_para('3.3 Case 2')
if p:
    replace_in_para(p, '3.3 Case 2', '3.2 Case 2')

# Bayley-III scores for Case 3 (per user: 75 / 65 / 79)
replace_all_in_doc('(Cognitive 65, Language 56, Motor 61)',
                   '(Cognitive 75, Language 65, Motor 79)')

# Note: the same string appears twice in the original (Case 1 + Case 3 typo).
# We need only the Case 3 one fixed; the Case 1 one stays at 65/56/61.
# But replace_all_in_doc replaced ALL — including Case 1. We need to undo Case 1.
# Re-find Case 1 paragraph and revert it.
p_case1 = find_para('The proband is a 4-year-old girl')
if p_case1 and '(Cognitive 75, Language 65, Motor 79)' in p_case1.text:
    replace_in_para(p_case1, '(Cognitive 75, Language 65, Motor 79)',
                              '(Cognitive 65, Language 56, Motor 61)')

# Recruitment dates / cohort source in §2.1
p21 = find_para('We enrolled three unrelated probands')
if p21:
    replace_in_para(p21,
        'Eligible patients were identified by their neurologists or clinical geneticists and recruited consecutively.',
        'Eligible probands were recruited consecutively at National Cheng Kung University Hospital between March 2023 and October 2024 (enrolment dates: 2 March 2023, 9 March 2023, and 21 October 2024) by referral from their attending neurologists or clinical geneticists.')

print('Step 1 (typos / simple fixes): done')

doc.save(OUTPUT)
print('Intermediate save:', OUTPUT)

# Re-load to make sure we work on the latest state if re-running incrementally.
# (When we run in one shot this is harmless.)

# ===========================================================================
# Step 2. Author / corresponding author block
# ===========================================================================
# Original block: para 8 "Corresponding Author:", 9 "Pao-Lin Kuo, ...", 10 "E-mail: paolinkuo@gmail.com"
# Replace with two corresponding authors: Yen-Yin Chou (lead) + Pao-Lin Kuo

p_corr_label = find_para('Corresponding Author')
if p_corr_label:
    replace_in_para(p_corr_label, 'Corresponding Author:', 'Corresponding authors:')

p_paolin = find_para('Pao-Lin Kuo, Department')
p_email = find_para('E-mail: paolinkuo')

if p_paolin:
    # Convert this line to: lead corresponding (Yen-Yin Chou)
    p_paolin.runs[0].text = ''
    for r in p_paolin.runs[1:]:
        r.text = ''
    p_paolin.runs[0].text = ('Yen-Yin Chou (lead corresponding author), Department of Genomic Medicine and '
                              'Department of Pediatrics, National Cheng Kung University Hospital, '
                              '138 Sheng-Li Road, Tainan 704, Taiwan. E-mail: yenyin01@gmail.com')

if p_email:
    p_email.runs[0].text = ('Pao-Lin Kuo, Department of Obstetrics and Gynecology, '
                             'National Cheng Kung University Hospital, '
                             '138 Sheng-Li Road, Tainan 704, Taiwan. E-mail: paolinkuo@gmail.com')
    for r in p_email.runs[1:]:
        r.text = ''

# ===========================================================================
# Step 3. Abstract — insert before "1. Introduction"
# ===========================================================================
p_intro = find_para('1. Introduction')
if p_intro:
    abstract_heading = insert_para_before(p_intro, 'Abstract', style='Heading 2')
    abstract_body = ('Apparently balanced chromosomal rearrangements identified by '
        'karyotyping are increasingly linked to neurodevelopmental disorders '
        '(NDDs), but their pathogenicity is often underestimated because '
        'cytogenetics and microarray cannot resolve cryptic imbalances or define '
        'breakpoint context. We applied integrated long-read sequencing (LRS) and '
        'optical genome mapping (OGM) to three unrelated NDD probands with '
        'apparently balanced karyotypes. In Case 1, inv(8)(p23p11.2) was refined '
        'into a deletion–inversion–deletion event within the 8p23.1 microdeletion '
        'region, with breakpoints implicating REPD/REPP-mediated non-allelic '
        'homologous recombination combined with microhomology-mediated end joining. '
        'In Case 2, inv(7)(q11.21q22) disrupted AUTS2 in intron 2, yielding a '
        'molecular diagnosis of AUTS2-related syndrome that prior whole-exome '
        'sequencing (WES) had missed; retrospective Manta and Delly re-analysis of '
        'the WES BAM showed near-zero coverage at the breakpoint (0.07× versus '
        'median 106× across captured AUTS2 exons), confirming an assay-level rather '
        'than analytical limitation. In Case 3, two apparently balanced '
        'translocations were resolved into a multi-chromosomal complex genomic '
        'rearrangement involving chromosomes 2, 5, 7 and 13, with CNTNAP2 disruption '
        'plausibly contributing to the autism/ADHD phenotype. Combining LRS and OGM '
        'converts apparently balanced karyotypes into precise molecular '
        'characterisations and exposes cryptic pathogenic events that escape '
        'conventional testing.')
    insert_para_after(abstract_heading, abstract_body, style='Normal')

# ===========================================================================
# Step 4. §2.5.2 / §2.6.2 — methods additions (HiFiCNV + SNV pipeline)
# ===========================================================================
# Methods §2.6.2 paragraph 2 (about pbmm2/pbsv/DeepVariant): add HiFiCNV
p_align = find_para('For alignment, HiFi reads were mapped')
if p_align:
    # Insert HiFiCNV mention after pbsv calling sentence
    replace_in_para(p_align,
        'SVs were called with pbsv (v2.10.0) using default analysis parameters.',
        'SVs were called with pbsv (v2.10.0) using default analysis parameters. '
        'Copy-number variants were called with HiFiCNV (v1.0.1) using default '
        'parameters and the recommended GRCh38 exclude-region BED supplied with '
        'the tool, with sex automatically inferred from chromosome X / Y depth.')

# Methods §2.6.2 paragraph 3 (downstream interpretation): expand SNV pipeline
p_downstream = find_para('For downstream interpretation, CNVs and SVs')
if p_downstream:
    p_downstream.runs[0].text = ''
    for r in p_downstream.runs[1:]:
        r.text = ''
    p_downstream.runs[0].text = (
        'For downstream interpretation, CNVs and SVs were annotated and prioritised '
        'using AnnotSV (https://www.lbgi.fr/AnnotSV/), while SNVs and indels were '
        'annotated with the latest releases of Ensembl Variant Effect Predictor '
        '(VEP) and ANNOVAR. Population allele frequencies were obtained from gnomAD '
        'v4.1, and clinical assertions from ClinVar (release 2025-12-08). Variants '
        'with allele frequency ≥0.01 in gnomAD (any population) were excluded as '
        'common. Functional impact was further assessed using AlphaMissense for '
        'missense variants, MetaRNN as a meta-predictor of pathogenicity, and '
        'SpliceAI for predicted splicing effects. The remaining variants were '
        'classified following the ACMG/AMP guideline. Breakpoints and copy-number '
        'changes were manually reviewed in IGV, and structural variant architecture '
        'and read-level configurations were further assessed using Ribbon '
        '(https://v2.genomeribbon.com/). The filtering and prioritisation workflows '
        'for HiFiCNV-derived CNVs and pbsv-derived SVs are summarised in '
        'Supplementary Tables 4 and 5, respectively.')

print('Step 2-4 (authors / abstract / methods): done')
doc.save(OUTPUT)
print('Saved:', OUTPUT)

# ===========================================================================
# Step 5. Results §3.1 — note about Case 1 OGM CNV bias
# ===========================================================================
p_ogm_case1 = find_para('OGM, however, identified fusion calls linking 8p23.1')
if p_ogm_case1:
    # Append a sentence at the end of this paragraph
    if 'OGM CNV counts in Case 1 were elevated' not in p_ogm_case1.text:
        # Add to last run preserving formatting
        extra = (' We note that the genome-wide OGM CNV count in Case 1 was '
                 'elevated owing to systematic coverage bias (Supplementary '
                 'Table 1); the two true 8p deletions described above were '
                 'nevertheless concordantly detected by OGM, LRS, and SNP '
                 'array (Figure 1B–E), so the conclusions in Case 1 are not '
                 'affected by this bias.')
        # Append to last run
        p_ogm_case1.runs[-1].text = p_ogm_case1.runs[-1].text + extra

# ===========================================================================
# Step 6. Figure 1G dual reference within same paragraph
# ===========================================================================
# Original passage: "...REPP and REPD as the source of the inversion (Figure 1G).
# Sequence inspection of the remaining two junctions ...consistent with
# microhomology-mediated end-joining (MMEJ)(Figure 1G)."
# Merge to a single (Figure 1G) at end.
p_mech = find_para('Breakpoint-level analysis provided insight')
if p_mech:
    replace_in_para(p_mech,
        'as the source of the inversion (Figure 1G).',
        'as the source of the inversion.')
    replace_in_para(p_mech,
        'compatible with microhomology-mediated end-joining (MMEJ)(Figure 1G).',
        'compatible with microhomology-mediated end-joining (MMEJ); the proposed mechanism is summarised in Figure 1G.')

# ===========================================================================
# Step 7. Discussion — Case 3 chromoanagenesis softening
# ===========================================================================
# Find the Discussion paragraph for Case 3 and replace the chromoanagenesis sentence
p_disc_case3 = find_para_contains('two de novo apparently balanced translocations')
if p_disc_case3:
    replace_in_para(p_disc_case3,
        'The shattering and reassembly of segments on multiple chromosomes is reminiscent of chromothripsis or chromoplexy, types of chromoanagenesis characterized by multiple double-strand breaks occurring simultaneously and rejoining in abnormal configurations [28, 31, 32].',
        'The shattering and reassembly of segments on multiple chromosomes is highly reminiscent of chromoanagenesis — a class of complex genomic rearrangement (CGR) that includes chromothripsis and chromoplexy and is thought to arise from multiple double-strand breaks resolving simultaneously into abnormal configurations [28, 30, 31, 32]. Our case fulfils several qualitative hallmarks of chromoanagenesis (clustered breakpoints distributed across two reciprocal exchanges, focal copy-number losses, and inverted orientations within rearranged segments); a formal classification under the chromothripsis criteria of Korbel and Campbell [REF_KORBEL] would require additional copy-number oscillation analysis. Importantly, the simultaneous detection of these multiple breakpoints required integrated LRS and OGM, since each platform alone would have under-resolved the rearrangement.')

# ===========================================================================
# Step 8. Discussion — Case 3 CNTNAP2 vs 7q deletion contribution
# ===========================================================================
# Append a sentence to the Case 3 discussion paragraph about deletion content
if p_disc_case3:
    extra_case3 = (' We acknowledge that the patient harbours both a CNTNAP2-disrupting '
                   'breakpoint and a co-occurring ~5.4 Mb 7q34–q35 deletion (Figure 3E), '
                   'which encompasses additional OMIM-listed genes (BRAF, AGK, CLCN1, '
                   'TBXAS1, WEE2, SSBP1, TAS2R38, PRSS1, PRSS2, TRPV6, KEL, CASP2, '
                   'NOBOX, TPK1, SLC37A3); with current data we cannot fully partition '
                   'the contribution of CNTNAP2 disruption versus the contiguous-gene '
                   'deletion to the proband\'s neurodevelopmental phenotype, although '
                   'CNTNAP2 has the strongest a priori evidence for autism-spectrum '
                   'and language phenotypes [23-26].')
    p_disc_case3.runs[-1].text = p_disc_case3.runs[-1].text + extra_case3

# ===========================================================================
# Step 9. Discussion — Case 2 retrospective Manta + Delly re-analysis
# ===========================================================================
p_disc_case2 = find_para_contains('The second case carried a de novo inversion on chromosome 7')
if p_disc_case2:
    replace_in_para(p_disc_case2,
        'This outcome exemplifies how genome-wide structural resolution can reveal pathogenic rearrangements disrupting known disease genes at intronic or intergenic breakpoints that are difficult to reconstruct accurately with WES [15].',
        'This outcome exemplifies how genome-wide structural resolution can reveal pathogenic rearrangements disrupting known disease genes at intronic or intergenic breakpoints that are difficult to reconstruct accurately with WES [15]. To verify that the failure of the prior WES was a fundamental assay limitation rather than an artefact of the analysis pipeline, we retrospectively re-analysed the original WES BAM with two methodologically distinct short-read SV callers — Manta v1.6.0 [REF_MANTA] (run with --exome and without restricting --callRegions) and Delly v1.2.6 [REF_DELLY] (default parameters) — using the GRCh38 + hs38d1 + HLA reference. Both callers operated normally on chromosome 7 (Manta: 129 candidate SVs; Delly: 142 SV records), yet neither returned any record within a ±50 kb window of either LRS-defined breakpoint, including in Manta\'s most permissive candidateSV output (Supplementary Table 6). Per-base read depth at the two breakpoints was 0.07× and 0.16×, against 154.81× at a captured positive-control exon (TP53 exon 4) and a median of 106.09× across the 19 captured AUTS2 exons (range 35.90×–274.85×; NM_015570.4) — a more than 1,500-fold deficit relative to the captured AUTS2 exonic baseline. The concordant negative result across two independent callers, together with the near-zero local read depth, demonstrates that the AUTS2 intron 2 breakpoint and its 7q21.3 partner lie outside exome capture territory and are inaccessible to short-read WES at the read level, regardless of the variant caller applied.')

# ===========================================================================
# Step 10. Discussion — comparative table 2 vs table 3 sensitivity note
# ===========================================================================
p_compare = find_para_contains('comparative performance of LRS and OGM')
if p_compare:
    extra_cmp = (' We further note that the apparent disparity in raw SV counts '
                 'between OGM (Table 2) and LRS (Table 3) reflects different size '
                 'sensitivities and filtering stages rather than a real difference '
                 'in detection breadth: OGM has a lower size limit of approximately '
                 '500 bp and the counts in Table 2 are already population-rarity '
                 'filtered against the Bionano control database, whereas LRS pbsv '
                 'calls were enumerated at SVLEN ≥50 bp prior to population '
                 'filtering. After equivalent rarity filtering, the two platforms '
                 'returned comparable numbers of high-confidence rare SVs per case '
                 '(Table 2 versus the population-filtered SV row of Table 3).')
    p_compare.runs[-1].text = p_compare.runs[-1].text + extra_cmp

print('Step 5-10 (results / discussion edits): done')
doc.save(OUTPUT)
print('Saved:', OUTPUT)

# ===========================================================================
# Step 11. Discussion — limitations: add Sanger/trio caveat
# ===========================================================================
p_lim = find_para('Our study has several limitations.')
if p_lim:
    extra_lim = (' Fifth, breakpoint junctions and parental trio samples could not be '
                 'confirmed by Sanger sequencing because the families declined '
                 'additional blood sampling for personal reasons; we therefore '
                 'relied on the orthogonal concordance of LRS and OGM (and SNP '
                 'array, where applicable) as independent lines of evidence '
                 'supporting the breakpoints and copy-number changes reported '
                 'here. Future cohorts incorporating routine trio confirmation '
                 'would further strengthen inference of de novo origin and would '
                 'exclude rare instances of low-level parental mosaicism.')
    # append to last run of paragraph
    p_lim.runs[-1].text = p_lim.runs[-1].text + extra_lim

# ===========================================================================
# Step 12. Conclusion + abstract softening for Case 1 (Major #9 upper)
# ===========================================================================
# Look for "conclusive molecular diagnosis" or similar in summary paragraph
replace_all_in_doc(
    'converting ambiguous karyotypic findings into definitive molecular diagnoses',
    'refining ambiguous karyotypic findings into precise molecular characterisations'
)
# In summary "supporting previous calls" sentence
replace_all_in_doc(
    'transform an initially "balanced" finding into a conclusive molecular diagnosis',
    'transform an initially "balanced" finding into a precise molecular characterisation'
)
# Discussion implications paragraph: soften "definitive molecular diagnosis"
replace_all_in_doc(
    'providing definitive molecular diagnosis, and ending the diagnostic odyssey',
    'providing a refined molecular characterisation that resolves the apparent karyotype as an unbalanced deletion-inversion-deletion event in Case 1 and identifies dosage-sensitive disease genes (the 8p23.1 microdeletion region, AUTS2, CNTNAP2) as the most likely contributors to each proband\'s phenotype'
)

# ===========================================================================
# Step 13. Insert end-of-paper sections after Discussion, before References
# ===========================================================================
# Find "Reference:" header
p_ref_header = find_para('Reference')
if p_ref_header:
    # Insert sections in REVERSE order using insert_para_before so they end up
    # in the correct order: Ethics → Acknowledgements → Author contributions →
    # Funding → Competing interests → Data availability → Code availability → References
    sections = [
        # (heading, body)
        ('Ethics declarations',
         'The study was approved by the Institutional Review Board of National '
         'Cheng Kung University Hospital, Taiwan (IRB No. A-BR-109-045). Written '
         'informed consent was obtained from all adult participants or the legal '
         'guardians of paediatric probands; assent was obtained from children '
         'when developmentally appropriate. All participants (or legal guardians) '
         'provided written informed consent for the use of their de-identified '
         'medical and genetic data, including for publication.'),

        ('Acknowledgements',
         'We thank Wilson Cheng (PacBio) for providing bioinformatics analysis '
         'support, Blossom Biotechnologies, Inc. (Taiwan) for assistance in '
         'coordinating the long-read sequencing workflow, and Pharmigene, Inc. '
         '(Taiwan) for assistance with optical genome mapping.'),

        ('Author contributions',
         'Y.-M.C. — Methodology, Data curation, Formal analysis, Investigation, '
         'Software, Visualization, Writing – original draft. Y.-W.P. — '
         'Investigation, Resources. Y.-Y.C. — Investigation, Resources, '
         'Supervision, Writing – review & editing. M.-C.T. — Supervision, '
         'Writing – review & editing. P.-M.W. — Resources, Writing – review & '
         'editing. P.-L.K. — Conceptualization, Funding acquisition, Methodology, '
         'Supervision, Writing – review & editing.'),

        ('Funding',
         'This work was supported by the Clinical Medical Research Center, '
         'National Cheng Kung University Hospital (grant number NCKUH-11009023). '
         'The funder had no role in study design, data collection and analysis, '
         'decision to publish, or preparation of the manuscript.'),

        ('Competing interests',
         'The authors declare no competing interests.'),

        ('Data availability',
         'The data that support the findings of this study are available from '
         'the corresponding authors upon request. Restrictions apply because the '
         'data contain information that could compromise the privacy of the '
         'participants; access is therefore subject to local ethics-board approval.'),

        ('Code availability',
         'Custom analysis scripts (including the retrospective Manta and Delly '
         're-analysis pipelines for Case 2 and the AUTS2 per-exon coverage '
         'computation) are available from the corresponding authors upon request.'),
    ]

    # Insert in REVERSE order so they appear in the listed order before
    # the References header.
    anchor = p_ref_header
    for heading, body in reversed(sections):
        body_para = insert_para_before(anchor, body, style='Normal')
        head_para = insert_para_before(body_para, heading, style='Heading 2')
        anchor = head_para

print('Step 11-13 (limitation / softening / end sections): done')
doc.save(OUTPUT)
print('Saved:', OUTPUT)

# ===========================================================================
# Step 14. Add new references and resolve [REF_*] placeholders
# ===========================================================================
# Existing references go up to [41]. We add three new refs:
#   [42] Korbel JO, Campbell PJ. Cell. 2013 (chromothripsis criteria)
#   [43] Chen X, Saunders CT et al. Manta. Bioinformatics 2016
#   [44] Rausch T et al. Delly. Bioinformatics 2012

# Find last existing reference (e.g. "[41]")
last_ref = None
for p in doc.paragraphs:
    if p.text.startswith('[41]'):
        last_ref = p
        break

if last_ref:
    new_refs = [
        '[42] Korbel JO, Campbell PJ. Criteria for inference of chromothripsis in cancer genomes. Cell. 2013;152:1226-36.',
        '[43] Chen X, Schulz-Trieglaff O, Shaw R, Barnes B, Schlesinger F, Källberg M, et al. Manta: rapid detection of structural variants and indels for germline and cancer sequencing applications. Bioinformatics. 2016;32:1220-2.',
        '[44] Rausch T, Zichner T, Schlattl A, Stütz AM, Benes V, Korbel JO. DELLY: structural variant discovery by integrated paired-end and split-read analysis. Bioinformatics. 2012;28:i333-9.',
    ]
    anchor = last_ref
    for txt in new_refs:
        # Use Normal style; some EndNote bibliographies use a special style,
        # but Normal is safe and editable.
        anchor = insert_para_after(anchor, txt, style='Normal')

# Replace placeholders with the assigned numbers
replace_all_in_doc('[REF_KORBEL]', '[42]')
replace_all_in_doc('[REF_MANTA]',  '[43]')
replace_all_in_doc('[REF_DELLY]',  '[44]')

# ===========================================================================
# Step 15. Add Supplementary Table 6 stub for Manta + Delly + AUTS2 depth
# ===========================================================================
# Append at end of supplementary section
last_sup = None
for p in doc.paragraphs:
    if p.text.startswith('Supplementary Figure 1'):
        last_sup = p
        break

# Find the last paragraph in the document if Supplementary Figure 1 not detected
if last_sup is None:
    last_sup = doc.paragraphs[-1]

# Insert a new Sup Table 6 header after the entire current supplementary section
# (placed at the very end so the user can position it later)
last = doc.paragraphs[-1]
sup6_block = [
    ('Supplementary Table 6. Retrospective short-read SV re-analysis of the Case 2 WES BAM (Manta v1.6.0 and Delly v1.2.6) and per-base read depth at AUTS2 exons (NM_015570.4) and at the LRS-defined inversion breakpoints.', 'Normal'),
    ('Footnote. Manta was run with --exome and without restricting --callRegions; Delly was run with default parameters. The reference used was GRCh38_full_analysis_set_plus_decoy_hla.fa. Per-base mean depth was computed with samtools depth -a. Across the 19 captured AUTS2 exons, mean depth ranged from 35.90× to 274.85× (median 106.09×). The two LRS-defined breakpoints (chr7:69,938,691, AUTS2 intron 2; chr7:95,720,897, 7q21.3) showed mean depths of 0.07× and 0.16×, respectively, and 0 / 0 records within ±50 kb in Manta diploidSV / candidateSV and Delly. Genome-wide, Manta produced 129 chromosome-7 candidate SVs and Delly produced 142 chromosome-7 SV records, confirming that both callers operated normally.', 'Normal'),
]
for txt, style in sup6_block:
    last = insert_para_after(last, txt, style=style)

print('Step 14-15 (references / supplementary table 6): done')
doc.save(OUTPUT)
print('Saved:', OUTPUT)

# ===========================================================================
# Step 16. Fix typo in Table 3 (LRS SV summary) — Case 3 duplications "27,38" -> "2,738"
# ===========================================================================
# The cell content is split across three runs ('27' / ',' / '38') in the
# original docx, so per-run substring matching fails. Detect by cell.text
# and rewrite the runs of the first paragraph.
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if '27,38' in cell.text and '2,738' not in cell.text:
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].text = '2,738'
                    for r in p.runs[1:]:
                        r.text = ''

print('Step 16 (table 3 typo): done')
doc.save(OUTPUT)
print('Final save:', OUTPUT)

# ===========================================================================
# Verification
# ===========================================================================
out = docx.Document(OUTPUT)
print()
print('=== VERIFICATION ===')
print('Total paragraphs in revised:', len(out.paragraphs))

checks = [
    ('Abstract heading', lambda d: any(p.text.startswith('Abstract') for p in d.paragraphs)),
    ('3.2 Case 2 numbering', lambda d: any(p.text.startswith('3.2 Case 2') for p in d.paragraphs)),
    ('Bayley-III Case 3 75/65/79', lambda d: any('(Cognitive 75, Language 65, Motor 79)' in p.text for p in d.paragraphs)),
    ('Bayley-III Case 1 still 65/56/61', lambda d: any('(Cognitive 65, Language 56, Motor 61)' in p.text for p in d.paragraphs)),
    ('UHMW (UMWH fixed)', lambda d: any('UHMW gDNA' in p.text for p in d.paragraphs) and not any('UMWH' in p.text for p in d.paragraphs)),
    ('Recruitment dates', lambda d: any('between March 2023 and October 2024' in p.text for p in d.paragraphs)),
    ('Yen-Yin Chou lead corresponding', lambda d: any('Yen-Yin Chou (lead corresponding' in p.text for p in d.paragraphs)),
    ('HiFiCNV v1.0.1', lambda d: any('HiFiCNV (v1.0.1)' in p.text for p in d.paragraphs)),
    ('SNV pipeline VEP+ANNOVAR', lambda d: any('Ensembl Variant Effect Predictor' in p.text for p in d.paragraphs)),
    ('gnomAD v4.1', lambda d: any('gnomAD v4.1' in p.text for p in d.paragraphs)),
    ('AlphaMissense+SpliceAI', lambda d: any('AlphaMissense' in p.text and 'SpliceAI' in p.text for p in d.paragraphs)),
    ('Manta v1.6.0 in Discussion', lambda d: any('Manta v1.6.0' in p.text for p in d.paragraphs)),
    ('Delly v1.2.6 in Discussion', lambda d: any('Delly v1.2.6' in p.text for p in d.paragraphs)),
    ('Sanger/trio limitation', lambda d: any('breakpoint junctions and parental trio samples could not be' in p.text for p in d.paragraphs)),
    ('Chromoanagenesis softened (Korbel)', lambda d: any('Korbel and Campbell' in p.text for p in d.paragraphs)),
    ('OMIM gene list (BRAF, AGK, ...)', lambda d: any('BRAF, AGK, CLCN1' in p.text for p in d.paragraphs)),
    ('Acknowledgements section', lambda d: any(p.text.startswith('Acknowledgements') for p in d.paragraphs)),
    ('Funding section', lambda d: any(p.text.startswith('Funding') for p in d.paragraphs)),
    ('Author contributions', lambda d: any(p.text.startswith('Author contributions') for p in d.paragraphs)),
    ('Competing interests', lambda d: any(p.text.startswith('Competing interests') for p in d.paragraphs)),
    ('Data availability', lambda d: any(p.text.startswith('Data availability') for p in d.paragraphs)),
    ('Code availability', lambda d: any(p.text.startswith('Code availability') for p in d.paragraphs)),
    ('Ethics declarations', lambda d: any(p.text.startswith('Ethics declarations') for p in d.paragraphs)),
    ('Ref [42] Korbel', lambda d: any(p.text.startswith('[42] Korbel JO') for p in d.paragraphs)),
    ('Ref [43] Manta', lambda d: any(p.text.startswith('[43] Chen X') for p in d.paragraphs)),
    ('Ref [44] Delly', lambda d: any(p.text.startswith('[44] Rausch T') for p in d.paragraphs)),
    ('Sup Table 6', lambda d: any(p.text.startswith('Supplementary Table 6') for p in d.paragraphs)),
    ('Pharmigene OGM', lambda d: any('Pharmigene' in p.text for p in d.paragraphs)),
    ('Blossom acknowledgement', lambda d: any('Blossom Biotechnologies' in p.text for p in d.paragraphs)),
    ('Wilson Cheng', lambda d: any('Wilson Cheng' in p.text for p in d.paragraphs)),
    ('Whereas → whereas', lambda d: not any('Moreover, Whereas' in p.text for p in d.paragraphs)),
    ('upper track (tract → track)', lambda d: not any('upper tract' in p.text for p in d.paragraphs)),
    ('Figure 3D coordinate (chr7:139)', lambda d: any('chr7:139 Mb' in p.text for p in d.paragraphs)),
    ('Table 3 Case 3 duplications fixed (no 27,38)', lambda d: not any('27,38' in (cell.text for tbl in d.tables for row in tbl.rows for cell in row.cells))),
]
for name, check in checks:
    try:
        ok = check(out)
        status = 'OK' if ok else 'FAIL'
    except Exception as e:
        status = f'ERR ({e})'
    print(f'  [{status}] {name}')

# ===========================================================================
# Step 17. Reference list -> Nature/npj Genomic Medicine style
# ===========================================================================
# Source format:
#   "[1] Warburton D. De novo balanced ...  Am J Hum Genet. 1991;49:995-1013."
# Target format (Nature-style; npj GM):
#   "1. Warburton, D. De novo balanced ...  Am J Hum Genet. 49, 995-1013 (1991)."
#   with journal italic, volume bold.
#
# In text:  "[1]" / "[1, 2]" / "[1-3]"  ->  superscript "1" / "1,2" / "1-3"

import re
from docx.shared import Pt

REF_RE = re.compile(
    r'^\[(\d+)\]\s+'
    r'(?P<authors>.+?)\.\s+'           # author list ending in ". "
    r'(?P<title>.+?)\.\s+'             # title ending in ". "
    r'(?P<journal>[^.]+)\.\s+'          # journal name ending in ". "
    r'(?P<year>\d{4})\s*;\s*'           # year;
    r'(?P<vol>\d+[A-Za-z]*)'            # volume
    r'(?:\((?P<issue>[^)]+)\))?'        # optional (issue)
    r'\s*:\s*'
    r'(?P<pages>[^.]+)\.?\s*$'
)

def reformat_authors(authors):
    """Convert 'Warburton D' style to 'Warburton, D.' style.
    Multiple authors separated by ', ' in the source."""
    out = []
    for a in [s.strip() for s in authors.split(',')]:
        if not a:
            continue
        if a.lower() == 'et al':
            out.append('et al.')
            continue
        # Format: "Lastname Initial(s)".  e.g., "Talkowski ME"
        # Take last token as initials, rest as surname.
        parts = a.rsplit(' ', 1)
        if len(parts) == 2 and re.fullmatch(r'[A-Z]{1,5}', parts[1]):
            initials = '.'.join(list(parts[1])) + '.'
            out.append(f'{parts[0]}, {initials}')
        else:
            out.append(a)
    if not out:
        return authors
    if len(out) == 1:
        return out[0]
    if 'et al.' in out:
        # combine: "First, A. et al."
        idx = out.index('et al.')
        # Standard: list first 5 authors then et al.; or just first + et al.
        # Source already has "et al" appended; keep first 5.
        kept = out[:5]
        if 'et al.' not in kept:
            return ', '.join(kept) + ' et al.'
        return ', '.join(kept[:-1]) + ' et al.'
    if len(out) == 2:
        return f'{out[0]} & {out[1]}'
    # 3+ authors, no et al.: comma-join with ampersand before last
    return ', '.join(out[:-1]) + ' & ' + out[-1]

def rewrite_reference_paragraph(para):
    text = para.text.strip()
    m = REF_RE.match(text)
    if not m:
        return False
    n = int(m.group(1))
    authors = reformat_authors(m['authors'])
    title = m['title']
    journal = m['journal']
    year = m['year']
    vol = m['vol']
    pages = m['pages'].replace('-', '–')  # en-dash for page ranges

    # Clear paragraph
    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    # Build runs: "{n}. {authors}. {title}. " (plain) +
    #             "{journal}" (italic) + " " (plain) +
    #             "{vol}" (bold) + ", {pages} ({year})." (plain)
    r1 = para.add_run(f'{n}. {authors}. {title}. ')
    r2 = para.add_run(journal); r2.italic = True
    r3 = para.add_run(' ')
    r4 = para.add_run(vol); r4.bold = True
    r5 = para.add_run(f', {pages} ({year}).')
    return True

# Apply to every paragraph that looks like a reference
ref_count = 0
for p in doc.paragraphs:
    if rewrite_reference_paragraph(p):
        ref_count += 1
print(f'Step 17a: {ref_count} reference paragraphs rewritten')

# In-text citations: convert "[N]" / "[N, M]" / "[N-M]" / "[N, M, ...]" to
# superscript without brackets.
CITE_RE = re.compile(r'\[\s*(\d+(?:\s*[-,]\s*\d+)*)\s*\]')

def superscriptify_citations(para):
    """Convert in-text bracket citations in this paragraph to superscript.
    Skips paragraphs that look like reference list entries."""
    if re.match(r'^\d+\.\s', para.text):
        return False
    if not CITE_RE.search(para.text):
        return False
    # Single-run case: easy to handle.
    if len(para.runs) == 1:
        return _superscriptify_single_run(para)
    # Multi-run: scan each run; if a citation is fully inside a run, split that run.
    runs_to_replace = []
    for run in list(para.runs):
        if CITE_RE.search(run.text):
            runs_to_replace.append(run)
    if not runs_to_replace:
        return False
    for run in runs_to_replace:
        _split_run_for_citations(para, run)
    return True

def _superscriptify_single_run(para):
    run = para.runs[0]
    txt = run.text
    matches = list(CITE_RE.finditer(txt))
    if not matches:
        return False
    # Need to rebuild paragraph runs preserving original formatting attributes
    base_font = run.font
    bold = run.bold; italic = run.italic
    name = base_font.name; size = base_font.size
    # Clear runs
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    last_end = 0
    def _new_run(text, sup=False):
        nr = para.add_run(text)
        if bold:   nr.bold = True
        if italic: nr.italic = True
        if name:   nr.font.name = name
        if size:   nr.font.size = size
        if sup:    nr.font.superscript = True
        return nr
    for m in matches:
        if m.start() > last_end:
            _new_run(txt[last_end:m.start()])
        # Strip the brackets and whitespace inside
        inner = m.group(1).replace(' ', '')
        _new_run(inner, sup=True)
        last_end = m.end()
    if last_end < len(txt):
        _new_run(txt[last_end:])
    return True

def _split_run_for_citations(para, run):
    txt = run.text
    if not CITE_RE.search(txt):
        return
    # Insert new runs by splitting around matches at the XML level.
    matches = list(CITE_RE.finditer(txt))
    bold = run.bold; italic = run.italic
    name = run.font.name; size = run.font.size
    parent = run._element.getparent()
    idx = list(parent).index(run._element)
    parent.remove(run._element)

    def _make_run_xml(text, sup=False):
        # use python-docx convenience by adding a temporary run, then moving
        tmp = para.add_run(text)
        if bold:   tmp.bold = True
        if italic: tmp.italic = True
        if name:   tmp.font.name = name
        if size:   tmp.font.size = size
        if sup:    tmp.font.superscript = True
        elem = tmp._element
        elem.getparent().remove(elem)
        return elem

    new_elems = []
    last_end = 0
    for m in matches:
        if m.start() > last_end:
            new_elems.append(_make_run_xml(txt[last_end:m.start()]))
        inner = m.group(1).replace(' ', '')
        new_elems.append(_make_run_xml(inner, sup=True))
        last_end = m.end()
    if last_end < len(txt):
        new_elems.append(_make_run_xml(txt[last_end:]))
    for el in new_elems:
        parent.insert(idx, el)
        idx += 1

cite_count = 0
for p in doc.paragraphs:
    if superscriptify_citations(p):
        cite_count += 1
print(f'Step 17b: {cite_count} paragraphs had citations superscripted')

doc.save(OUTPUT)
print('Saved after Step 17:', OUTPUT)

# ===========================================================================
# Step 18. Unify font for all body paragraphs
# ===========================================================================
# Use Times New Roman 12pt for body, keep heading styles' weights but normalise
# their typeface too.
from docx.shared import Pt

BODY_FONT = 'Times New Roman'
BODY_SIZE = Pt(12)
HEADING_SIZES = {'Heading 1': Pt(16), 'Heading 2': Pt(14),
                 'Heading 3': Pt(13), 'Heading 4': Pt(12)}

for p in doc.paragraphs:
    style_name = p.style.name if p.style else ''
    target_size = HEADING_SIZES.get(style_name, BODY_SIZE)
    for run in p.runs:
        # Don't reset bold/italic/superscript — preserve those
        run.font.name = BODY_FONT
        # Force eastAsia font name too (fixes mixed CJK)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), BODY_FONT)
        rfonts.set(qn('w:hAnsi'), BODY_FONT)
        rfonts.set(qn('w:cs'), BODY_FONT)
        rfonts.set(qn('w:eastAsia'), BODY_FONT)
        run.font.size = target_size

# Also normalise font for table cells
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = BODY_FONT
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.find(qn('w:rFonts'))
                    if rfonts is None:
                        from docx.oxml import OxmlElement
                        rfonts = OxmlElement('w:rFonts')
                        rpr.insert(0, rfonts)
                    rfonts.set(qn('w:ascii'), BODY_FONT)
                    rfonts.set(qn('w:hAnsi'), BODY_FONT)
                    rfonts.set(qn('w:cs'), BODY_FONT)
                    rfonts.set(qn('w:eastAsia'), BODY_FONT)
                    run.font.size = BODY_SIZE

print('Step 18 (font unification): done')
doc.save(OUTPUT)
print('Saved:', OUTPUT)

# ===========================================================================
# Step 19. Build a real Word table for Supplementary Table 6
# ===========================================================================
# Find the Sup Table 6 caption paragraph (we inserted it earlier)
sup6_caption = None
for p in doc.paragraphs:
    if p.text.startswith('Supplementary Table 6.'):
        sup6_caption = p
        break

if sup6_caption is not None:
    # Find the footnote paragraph (so we can insert table BETWEEN them)
    sup6_footnote = None
    cur = sup6_caption
    # Walk forward to find footnote
    for p in doc.paragraphs:
        if p is sup6_caption:
            keep_walking = True
            continue
        if p.text.startswith('Footnote.') and 'Manta' in p.text and 'Delly' in p.text:
            sup6_footnote = p
            break

    # Build table data
    header = ['Region (GRCh38)', 'Annotation', 'Length (bp)',
              'Mean depth (×)', 'Manta records (±50 kb)', 'Delly records (±50 kb)']
    rows = [
        ('chr17:7,675,995–7,676,272', 'TP53 exon 4 (positive control)', '278',  '154.81', 'n/a', 'n/a'),
        ('chr7:69,598,475–69,599,962', 'AUTS2 exon 1', '1,488', '35.90', 'n/a', 'n/a'),
        ('chr7:69,899,286–69,899,498', 'AUTS2 exon 2', '213', '274.85', 'n/a', 'n/a'),
        ('chr7:70,118,132–70,118,233', 'AUTS2 exon 3', '102', '110.64', 'n/a', 'n/a'),
        ('chr7:70,134,536–70,134,571', 'AUTS2 exon 4', '36', '160.83', 'n/a', 'n/a'),
        ('chr7:70,435,752–70,435,781', 'AUTS2 exon 5', '30', '103.77', 'n/a', 'n/a'),
        ('chr7:70,698,569–70,698,620', 'AUTS2 exon 6', '52', '100.29', 'n/a', 'n/a'),
        ('chr7:70,762,870–70,763,341', 'AUTS2 exon 7', '472', '195.63', 'n/a', 'n/a'),
        ('chr7:70,764,752–70,765,005', 'AUTS2 exon 8', '254', '105.28', 'n/a', 'n/a'),
        ('chr7:70,766,114–70,766,334', 'AUTS2 exon 9', '221', '228.38', 'n/a', 'n/a'),
        ('chr7:70,768,024–70,768,068', 'AUTS2 exon 10', '45', '63.27', 'n/a', 'n/a'),
        ('chr7:70,771,549–70,771,644', 'AUTS2 exon 11', '96', '106.09', 'n/a', 'n/a'),
        ('chr7:70,774,028–70,774,099', 'AUTS2 exon 12', '72', '165.08', 'n/a', 'n/a'),
        ('chr7:70,775,357–70,775,386', 'AUTS2 exon 13', '30', '123.83', 'n/a', 'n/a'),
        ('chr7:70,777,103–70,777,174', 'AUTS2 exon 14', '72', '91.78', 'n/a', 'n/a'),
        ('chr7:70,781,615–70,781,756', 'AUTS2 exon 15', '142', '236.77', 'n/a', 'n/a'),
        ('chr7:70,784,942–70,785,019', 'AUTS2 exon 16', '78', '91.26', 'n/a', 'n/a'),
        ('chr7:70,785,955–70,786,038', 'AUTS2 exon 17', '84', '95.80', 'n/a', 'n/a'),
        ('chr7:70,787,209–70,787,431', 'AUTS2 exon 18', '223', '168.65', 'n/a', 'n/a'),
        ('chr7:70,789,748–70,793,506', 'AUTS2 exon 19', '3,759', '64.74', 'n/a', 'n/a'),
        ('chr7:69,938,192–69,939,191', 'LRS breakpoint A (AUTS2 intron 2, ±500 bp)', '1,000', '0.07', '0', '0'),
        ('chr7:95,720,398–95,721,397', 'LRS breakpoint B (7q21.3, ±500 bp)',         '1,000', '0.16', '0', '0'),
        ('chr7 (whole chromosome)',     'Total candidate SV records (genome-wide context)', '–', '–',
            '129 (candidateSV)', '142'),
    ]

    nrows = len(rows) + 1  # +1 header
    ncols = len(header)

    # Create table at end of doc, then move it to right after caption
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    # Fill header
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        c.text = ''  # clear default
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
    # Fill data
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = tbl.rows[i].cells[j]
            c.text = ''
            run = c.paragraphs[0].add_run(val)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
    # Move table to right after caption paragraph
    sup6_caption._element.addnext(tbl._element)

print('Step 19 (Sup Table 6 actual table): done')
doc.save(OUTPUT)
print('Final saved:', OUTPUT)

# ===========================================================================
# Step 17c. Cleanup: remove '..' double period after author block
# (caused by author initials ending with '.' + our ". " separator)
# ===========================================================================
import re as _re
for p in doc.paragraphs:
    if not _re.match(r'^\d+\.\s', p.text):
        continue
    if any(p.text.startswith(pref) for pref in
           ('1. Introduction', '2. Materials', '3. Results', '4. Discussion')):
        continue
    if p.runs and '.. ' in p.runs[0].text:
        p.runs[0].text = p.runs[0].text.replace('.. ', '. ')

doc.save(OUTPUT)
print('Step 17c (double-period cleanup): done')
